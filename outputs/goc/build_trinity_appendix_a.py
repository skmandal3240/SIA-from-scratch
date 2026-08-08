#!/usr/bin/env python3
"""Build standalone Appendix A — Investor Due Diligence (Key Q&A) DOCX for Project TRINITY.

Maps to Investor_Due_Diligence_&_Valuation_Questions_for_Startups_2026.pdf categories.
Run: python3 outputs/goc/build_trinity_appendix_a.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_Appendix_A_Due_Diligence.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def q(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY


def a(text):
    doc.add_paragraph(text)


def num(text):
    p = doc.add_paragraph()
    r = p.add_run("Numeric example: " + text)
    r.italic = True
    r.font.color.rgb = GREY


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Appendix A — Investor Due Diligence (Key Questions & Answers)\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(14)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Working draft | August 2026 | SIA / Mandal Holdings\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
doc.add_paragraph("Answers map to the question categories in Investor Due Diligence & Valuation Questions for Startups (2026). Numeric figures are illustrative planning estimates unless marked as verified.")

# ============ 1. COMPANY OVERVIEW ============
h1("1. Company Overview")

q("Q1. What is the company's mission, history, stage, and legal status?")
a("Project TRINITY is the vehicle for building India's complete sovereign AI stack — compute, data, models, agents/OS, and applications — delivered through SIA, a family of on-device and cloud AI models. The company is being incorporated as a Private Limited Company in India (SPICe+ filing in progress, registered office recommended in Patna, Bihar for the Bihar Startup Policy; founder domicile Madhubani, Bihar). Status: pre-revenue, Seed/pre-Series-A stage. History: from-scratch SIA framework built and trained (nano transformer on CPU), LoRA fine-tune pipeline ready, investor deck and grant dossiers drafted, 5-pillar plan documented.")
num("Incorporation target: Q3 2026. Founder holds 100% initially. Headcount: 1 founder + 1–2 hires planned in MVP phase.")

q("Q2. Who are the founders and core team?")
a("Founder Saurabh Mandal — full-stack builder and AI system architect who built the entire from-scratch SIA stack (tokenizer, transformer, vision/audio/code encoders, diffusion heads, tool-calling, swarm scaffold) — ~2,000+ lines of original code — plus three prior products (ASTRO, ALICE, SIA). Domain coverage: model architecture, training, multimodal ML, product engineering, business planning, and the government-grant pipeline (TIDE 2.0, SISFS, IndiaAI, Bihar Startup Policy).")
num("1 founder; hiring 1–2 junior engineers in year 1; ML + infra hires years 2–3.")

q("Q3. What is the business model and how does the company create value?")
a("Multi-layer revenue: (1) SIA API usage (cloud tier), (2) enterprise subscriptions (SIA Pro + agents), (3) GPU cloud after the private cluster comes online, (4) managed AI services and licensing, (5) SDK/embed fees via SIA Studio. Value creation: a complete sovereign stack — every layer owns margin, and the from-scratch approach means no API-wrapper fees. Privacy is the product: on-device inference keeps data local, aligning with DPDP 2023.")
num("MVP phase targets: 3–5 enterprise pilots; ARR ₹0→₹25 lakh in year 2 (illustrative).")

q("Q4. What is the TAM, and what share can you capture?")
a("TAM: global AI infrastructure + applications market (multi-trillion USD). SAM: India AI market — enterprise, government, and consumer — estimated $25B+ by 2030. SOM: on-device Indian-language AI, a segment with no credible full-stack Indian player today; initial capture via enterprise pilots and government contracts.")
num("SOM plan: 1% of Indian AI spend by year 5 (illustrative; detailed model in financial appendices).")

q("Q5. Who are competitors and what is the competitive advantage?")
a("Global: OpenAI, Anthropic, Google — cloud-locked, English-first, no sovereign-India positioning. China: full sovereign stack exists; India has no equivalent. India: no credible full-stack on-device + sovereign player. Advantage: from-scratch IP (not a wrapper), on-device privacy, Indian language/context data, staged silicon roadmap (DLI/PLI), and the sovereign-infrastructure narrative for government procurement.")
num("Moats: patents/trade secrets on the from-scratch stack; proprietary Indian-language data; government alignment.")

# ============ 2. MARKET & CUSTOMERS ============
h1("2. Market & Customers")

q("Q6. Who are target customers and what traction exists?")
a("Tier 1: Indian government (sovereign AI contracts, IndiaAI-aligned). Tier 2: Indian enterprises needing on-device/private AI. Tier 3: developers embedding SIA via SDK. Traction: working from-scratch model (9-demo gauntlet passing: text, code, vision, audio-listen, audio-gen, image, video, tools, quantize), audio→language loop live, LoRA pipeline ready, grant dossiers drafted.")
num("Pilots: 0 signed (pre-revenue); 5 target in year 1 (illustrative).")

q("Q7. What is the customer acquisition strategy and unit economics?")
a("GTM sequence: developers → startups → SMEs → enterprises → government. Developer-first via SIA Studio/SDK builds organic adoption; government via sovereign narrative and grants; enterprise via direct sales and pilots. Unit economics (illustrative): CAC target ₹2–5 lakh per enterprise; LTV target ₹20–50 lakh over 4-year contracts; LTV/CAC target >3.")
num("LTV/CAC illustrative: LTV ₹30 lakh / CAC ₹5 lakh = 6.0× (healthy). Payback <12 months.")

# ============ 3. PRODUCT & TECHNOLOGY ============
h1("3. Product & Technology")

q("Q8. What problem does the product solve and what is its status?")
a("Problem: India lacks sovereign AI infrastructure; cloud AI is unusable offline for 900M+ users; data leaves devices to foreign servers. Solution: SIA — an AI operating system, a family of models (Nano 0.5–1B → Cloud 70B+) sharing one tokenizer, architecture, tool-calling interface, memory format, and runtime. Status: working from-scratch nano model trained on CPU (loss 2.78); vision VAE decoder added; audio listen/generate pipeline; LoRA fine-tune notebook ready.")
num("Demos passing: 9/9. Model params: nano (illustrative ~10–20M trainable). Training: 4,000 steps CPU.")

q("Q9. How is the product protected by intellectual property?")
a("All code owned under the company (founder-written, original). Planned: patent filings on the model architecture and training pipeline; trademarks for SIA; trade secrets for datasets and fine-tuning methodology; NDAs + IP-assignment for all hires.")
num("Target: 2–4 Indian patent filings in year 1–2; 1 trademark application for 'SIA'.")

q("Q10. What is the product roadmap and timeline?")
a("Phase 1 (2026–27): SIA Edge on-device companion MVP + first pilots. Phase 2 (2027–29): foundation models — Nano/Edge trained, Pro LoRA fine-tuned on Indian data. Phase 3 (2029–31): enterprise platform (agents, SDK, memory). Phase 4 (2031–35): national AI infrastructure (private cluster → hyperscale campus). Phase 5 (2035+): global expansion.")
num("Milestones: MVP demo Q4 2026; 5 pilots 2027; 1 government contract by 2028 (illustrative).")

# ============ 4. BUSINESS MODEL & STRATEGY ============
h1("4. Business Model & Strategy")

q("Q11. What is the revenue model and pricing strategy?")
a("Tiered: freemium consumer app (SIA Edge); usage-based API (SIA Cloud); per-seat enterprise subscriptions (SIA Pro); custom projects + managed services for government; SDK licensing for OEM/IoT. Pricing anchored on 'privacy as premium' and sovereignty.")
num("Illustrative: API ₹0.5–2/1K tokens; enterprise ₹5–20 lakh/yr; government contracts milestone-based.")

q("Q12. What are key unit economics? (CAC, LTV, churn, payback)")
a("Illustrative targets: gross margin 70–85% (software + on-device, no per-token cloud cost at edge); enterprise churn <5%/yr; CAC payback <12 months; LTV/CAC >3. Tracked from first revenue.")
num("LTV/CAC target 4–6×; net revenue retention target >110%.")

# ============ 5. FINANCIALS ============
h1("5. Financials (Historical & Projections)")

q("Q13. What are historical financial results?")
a("Pre-revenue to date. Spend to date: founder-funded compute (Colab/GCP), software licenses, filing fees. Books to be opened at incorporation; GAAP accounting from day one.")
num("Burn to date: <₹5 lakh (illustrative, founder-funded).")

q("Q14. What are financial projections and assumptions?")
a("Phase-gated, milestone-based — no hockey-stick claims. MVP (₹2–10 Cr): team + prototype + compute. Growth (₹25–100 Cr): foundation models. Scale (₹100–500 Cr): AI cloud. National (₹500–2,000 Cr): infrastructure. Global (₹2,000 Cr+): expansion. Each phase funds only after the prior milestone is hit. Assumptions: India AI market growth >25%/yr; government grant/contract pipeline; retention of IP in India.")
num("Year-3 revenue target: ₹10–50 Cr (illustrative). Cash runway after MVP ask: 18–24 months.")

# ============ 6. CAP TABLE & CAPITAL STRUCTURE ============
h1("6. Capital Structure & Valuation")

q("Q15. What is the current cap table?")
a("Pre-incorporation: 100% founder. Post-incorporation: founder holds 100% (₹1L paid-up / ₹10L authorised). ESOP pool to be carved (target 10–15%) at the first institutional round.")
num("Authorised ₹10,00,000; paid-up ₹1,00,000; 10,000 shares @ ₹10.")

q("Q16. What is the valuation strategy?")
a("Not a typical startup multiple. Position as infrastructure: asset-backed milestones (GPU cluster, models, patents, government contracts) + revenue, with sovereign-strategic premium. Valuation methods for reference: DCF (long-horizon), comparables (AI infra), precedent transactions, VC method — sensitivity table in financial appendices.")

# ============ 7. GOVERNANCE & LEGAL ============
h1("7. Governance, Legal & Tax")

q("Q17. What is the legal and tax structure?")
a("India Private Limited Company (SPICe+ INC-32). PAN/TAN auto-issued with CIN. GST registration when revenue begins. DPIIT/Startup India recognition (gates IndiaAI + SAMRIDH + TIDE 2.0). Bihar Startup Policy registration if Patna office. All contracts (employment, IP assignment, NDAs) in place from first hire.")
num("Incorporation cost ₹5–12k; 7–10 working days to CIN once founder docs submitted.")

q("Q18. What are the key risks?")
a("Capital intensity — mitigated by staged rent→build. Competition — differentiated by sovereignty + on-device privacy. Technology execution — mitigated by working from-scratch stack. Regulatory — DPDP alignment is an advantage. Talent — Indian dev pool + remote-first. AGI claims — strictly evidence-based, no overpromising.")

# ============ 8. CLOSING ============
h1("8. Ready-to-Answer Investor Questions (Quick Reference)")
for question, answer in [
    ("Why now?", "IndiaAI Mission funding 20 sovereign models; DPDP makes privacy an advantage; Jevons paradox favors cost leaders; no Indian full-stack sovereign player."),
    ("Why this team?", "Founder built the entire from-scratch stack alone — tokenizer, transformer, multimodal, tools — plus 3 prior products and a complete grant pipeline."),
    ("Why this market?", "India AI growing >25%/yr; 900M+ users offline; government demand for sovereign AI; no credible domestic competitor."),
    ("What is the competitive advantage?", "From-scratch IP, on-device privacy, Indian-language data, silicon roadmap, sovereign-government alignment."),
    ("How will capital be deployed?", "MVP ask ₹2–10 Cr: ~35% compute, ~30% team, ~20% data/fine-tuning, ~15% legal/SDK."),
    ("What milestones unlock the next round?", "MVP demo + 5 pilots + 1 government contract → Growth round (₹25–100 Cr) for foundation models."),
]:
    q(question)
    a(answer)

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")
