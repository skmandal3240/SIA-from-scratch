#!/usr/bin/env python3
"""Germany MSc AI/ML Roadmap DOCX — for Saurabh (Virar base, sister's place)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "Germany_MSc_AI_ML_Roadmap.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
s = doc.styles["Normal"]
s.font.name = "Calibri"
s.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY; r.font.size = Pt(16)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD; r.font.size = Pt(13)

def p(t): doc.add_paragraph(t)
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def note(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.color.rgb = GREY; r.font.size = Pt(9)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hh in enumerate(headers): t.rows[0].cells[i].text = hh
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row): t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()

# Cover
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("GERMANY MSc AI/ML — ROADMAP\n"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
r = doc.add_paragraph("Masters in Computer Science (AI & ML) · Winter 2026 / Summer 2027\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(13); r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Saurabh Mandal · base: Virar, Maharashtra (sister's place) · 12 August 2026\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(10); r.runs[0].font.color.rgb = GREY
note("Facts checked 12 Aug 2026. Requirements change yearly — verify on DAAD.de, uni-assist, and APS India before submitting. TODAY = 12 Aug 2026.")

# 1. Why Germany
h1("1. Why Germany (cost vs UK)")
table(["Item", "UK (Greenwich)", "Germany (public uni)"],
[
    ("Tuition", "£18,700/yr (~₹21L)", "€0 (semester fee ~€300 ≈ ₹30k)"),
    ("Living proof", "£~12k/yr", "Blocked account ~€11,904/yr (≈₹11L)"),
    ("Work", "20h/wk", "20h/wk (140 full days/yr)"),
    ("After degree", "2yr Graduate Route", "18mo job-seeker → EU Blue Card"),
    ("Language for admission", "English only", "Most AI MSc English-taught; German B1 for life/visa"),
])
p("Bottom line: Germany costs roughly ONE UK tuition year, total. And the AI/ML field in Germany is hiring aggressively.")

# 2. Two intake options
h1("2. Intake Options (as of 12 Aug 2026)")
table(["Intake", "Timing", "Status", "Verdict"],
[
    ("Winter 2026", "Oct 2026 — 6 weeks away", "Application windows mostly closed; visa lead time 6–12 wks impossible", "❌ Effectively missed — do not chase"),
    ("Summer 2027", "Apr 2027 — ~8 months away", "Applications open Sep–Nov 2026; deadlines Nov 2026–Jan 2027", "✅ THE realistic target"),
    ("Winter 2027", "Oct 2027 — 14 months away", "Backup if Summer 2027 slips", "Plan B"),
])
p("Recommendation: Summer 2027 is now the primary goal. Winter 2026 is unrealistic (visa alone needs 6–12 weeks after admission). Everything below is dated for Summer 2027.")

# 3. The 6 core requirements
h1("3. Core Requirements")
for i in [
    "APS certificate (German embassy verification of Indian degree) — 3–4 months lead time. APPLY NOW.",
    "German language: Goethe B1 minimum for visa; B2 for life/employability. Start at Goethe-Institut Mumbai or online.",
    "Degree: BE CSE 2025, CGPA 7.22 — meets typical 6.5–7.0+ requirement for AI MSc.",
    "English proof: IELTS/TOEFL if programme asks (many German unis accept CGPA + interview; check per-uni).",
    "Blocked account (Sperrkonto): ~€11,904 — build via Indian job + grants + savings.",
    "University applications via uni-assist / direct portals, usually Dec–Feb for winter, Jun–Sep for summer.",
]:
    bullet(i)

# 4. University shortlist (AI/ML, English-taught)
h1("4. University Shortlist (AI/ML English-taught)")
table(["University", "Programme", "Notes"],
[
    ("TU Munich (TUM)", "MSc Informatics / Robotics, Cognition, Intelligence", "Top-ranked; competitive; tuition-free state"),
    ("RWTH Aachen", "MSc Data Science / Computer Engineering", "Strong AI lab"),
    ("TU Berlin", "MSc Computer Science (AI track)", "Berlin tech scene"),
    ("KIT Karlsruhe", "MSc Informatics", "Excellent reputation"),
    ("University of Freiburg", "MSc Computer Science", "AI centre of excellence"),
    ("Saarland University", "MSc Computer Science (AI)", "MPI-INF connection"),
])
note("Each has specific English/GPA bars — apply to 4–6, 2 reach + 2 fit + 2 safe.")

# 5. German study plan (Virar base)
h1("5. German Study Plan (Virar, Maharashtra — from Aug 2026)")
table(["Milestone", "Target date", "How"],
[
    ("Goethe A1", "Dec 2026", "Online/Goethe Mumbai — ~3–4 months"),
    ("Goethe A2", "Apr 2027", "Intensive course + daily practice"),
    ("Goethe B1", "Jul 2027", "Exam — visa minimum (for Summer 2027 entry)"),
    ("Goethe B2 (optional)", "Dec 2027", "Stronger employability / winter 2027 backup"),
])
bullet("Daily: 1–2 hrs German (Duolingo + Goethe materials + German music/TV).")
bullet("Mumbai Goethe-Institut is 1–1.5h from Virar — weekend classes possible.")
bullet("Extra: join German speaking groups (Mumbai meetups) — free practice.")

# 6. Budget
h1("6. Budget (indicative, ₹)")
table(["Item", "Amount"],
[
    ("APS fee", "₹17,000–18,000"),
    ("Goethe exams (A1–B1)", "₹15,000–30,000"),
    ("Blocked account (€11,904)", "~₹11,00,000"),
    ("Flights + visa", "₹1,00,000"),
    ("Semester fees (first year)", "~₹30,000"),
    ("Total target savings", "~₹13,00,000"),
])
bullet("Funding routes: Indian job (Python/ML junior ₹4–12L/yr) + TIDE 2.0 grant (₹25L) + pilot income.")
bullet("Blocked account can be built in ~12 months with a job (₹60–80k/month saving).")

# 7. Timeline
h1("7. Month-by-Month Plan (12 Aug 2026 → Summer 2027)")
table(["When", "Action"],
[
    ("Aug 2026 (NOW)", "Apply APS immediately; book Goethe A1; open blocked-account research"),
    ("Sep–Oct 2026", "A1 classes; university shortlist final; watch summer-2027 portals open"),
    ("Nov 2026", "A1 exam; submit uni applications (first deadlines Nov 2026–Jan 2027)"),
    ("Dec 2026", "A2 starts; submit remaining applications"),
    ("Jan–Feb 2027", "A2 exam; admission letters start arriving"),
    ("Mar–Apr 2027", "B1 classes; accept admission; blocked account funding (€11,904)"),
    ("May–Jun 2027", "B1 exam; visa appointment (German embassy)"),
    ("Jul–Aug 2027", "Visa grant; flights; housing via uni dorms"),
    ("Sep 2027", "Pre-departure; blocked account activated"),
    ("Apr 2027 (intake)", "Summer semester starts — if all went well, this is the goal"),
])
note("If Summer 2027 slips → Winter 2027 becomes Plan B with B2 completed by Dec 2027.")

# 8. Keep SIA alive
h1("8. Keep SIA Alive Through It All")
bullet("SIA repo is your portfolio — German AI employers WILL check GitHub.")
bullet("Use free GPU tiers (Colab/Kaggle) to keep training while studying.")
bullet("Virar base = zero rent = money goes to German fund.")
bullet("Company (TRINITY) can stay Bengaluru-registered; run it remotely.")

# 9. Checklist
h1("9. Immediate Next Steps (this week)")
for i in ["Apply for APS on apsindia.org", "Book Goethe A1 course (online or Mumbai)", "Refresh resume + GitHub README (SIA as flagship)", "Open blocked-account research (Deutsche Bank/Fintiba)", "Shortlist 5 universities and note their deadlines"]:
    bullet("☐ " + i)

doc.save(OUT)
print(f"SAVED {OUT}")