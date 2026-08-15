#!/usr/bin/env python3
"""SARA Funding-Ready Package — PoC complete, legal checklist, pitch assets."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "SARA_Funding_Ready_Package.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
s = doc.styles["Normal"]
s.font.name = "Calibri"
s.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = NAVY; r.font.size = Pt(16)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = GOLD; r.font.size = Pt(13)

def p(t): doc.add_paragraph(t)
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def note(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.color.rgb = GREY; r.font.size = Pt(9)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hh in enumerate(headers): t.rows[0].cells[i].text = hh
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row): t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()

# Cover
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("SARA — FUNDING-READY PACKAGE\n"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
r = doc.add_paragraph("From-scratch multimodal transformer with agent runtime\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(13); r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Repo: https://github.com/skmandal3240/SARA · Date: 12 August 2026\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(10); r.runs[0].font.color.rgb = GREY
note("Code is PoC-complete. This package = legal/grant readiness + reviewer runbook.")

# 1. PoC Status
h1("1. PoC Status (Code — DONE)")
table(["Pillar", "Status", "Evidence"],
[
    ("SEE (vision)", "✅ Working", "Image encoder + cross-attn + caption generate; demos.py PASS"),
    ("TALK (audio)", "✅ Working", "Log-mel encoder + Griffin-Lim vocoder; wav I/O real"),
    ("CODE", "✅ Working", "AST sandbox + python_exec tool; agent writes/runs fib.py → 55"),
    ("CREATE IMAGE", "✅ Working", "Conv-transpose RGB 64×64; trained on synthetic shapes"),
    ("CREATE VIDEO", "✅ Working", "8-frame GIF decoder; low-res but pipeline runs"),
    ("CREATE SONG", "✅ Working", "SongHead + additive synth (melody/bass/pad/hats)"),
    ("TOOLS / AGENTS", "✅ Working", "Tool protocol, ReAct loop, planner, swarm; agent + swarm demos PASS"),
    ("EDGE RUNTIME", "✅ Working", "Profiles, INT8, paging, mesh (in-process); demos_edge.py PASS"),
    ("PRIVACY KERNEL", "✅ Working", "Grants, vault, audit, local learn hooks; tests PASS"),
    ("TESTS", "✅ 41/41 PASS", "pytest tests/ -q → 41 passed"),
    ("DEMO GAUNTLET", "✅ 9/9 PASS", "python demos.py → all 9 demos pass"),
])
p("Run yourself (CPU, no GPU, no cloud keys):")
p("```bash")
p("git clone https://github.com/skmandal3240/SARA && cd SARA")
p("python3 -m venv .venv && source .venv/bin/activate")
p("pip install -r requirements.txt")
p("python -m pytest tests/test_edge.py tests/test_privacy.py tests/test_data_catalog.py tests/test_tools.py -q")
p("python demos_edge.py  # must exit 0")
p("```")

# 2. What a reviewer sees
h1("2. What a Grant Reviewer Should See")
bullet("A from-scratch transformer (10.8M params) with GQA, RoPE, RMSNorm, SwiGLU — no API wrappers.")
bullet("Agent runtime that emits `<|tool_call|>` JSON, runs sandboxed Python, retries on error.")
bullet("Edge demos: CCTV see-path (cloud denied), Phone agent (cloud denied), 2-node mesh DAG.")
bullet("Privacy: grants/vault/audit/learn hooks — data never leaves device without consent.")
bullet("All runs on CPU in ~15 min train + 8 min test. No GPU required.")

# 3. Legal/Company Blockers (THE gap)
h1("3. Legal & Company Blockers — MUST FIX BEFORE GRANT SUBMISSION")
table(["Blocker", "Why It Gates Funding", "Owner", "Status"],
[
    ("Private Limited incorporation", "TIDE / IndiaAI / DLI require Indian entity", "Rahul (paperwork)", "❌ NOT DONE"),
    ("DPIIT Startup India recognition", "IndiaAI Mission gate; compute grants stall without it", "Rahul + Amit", "❌ NOT DONE"),
    ("≥51% Indian ownership", "IndiaAI / startup definition", "Founder + cap table", "❌ Company not formed"),
    ("In-house model ownership", "IndiaAI wants Indian-owned model, not wrapper", "Code is in-house", "✅ Code exists"),
    ("Pick ONE TIDE CoE", "Applications go via CoE, not generic inbox", "Founder", "❌ NOT PICKED"),
])
note("Without these, grant forms bounce. Incorporation + DPIIT = ~2–3 weeks with Rahul's help.")

# 4. Grant Targets & What to Show
h1("4. Grant Targets (from docs/GRANTS.md)")
table(["Instrument", "Lane", "What We Show", "Amount (PLAN.md)"],
[
    ("MeitY TIDE 2.0", "Infrastructure/transport (CCTV) OR Agriculture (drone)", "Edge CCTV/drone safety demos; SARA runtime; Pvt Ltd", "EiR ~₹4–7L; prototype up to ~₹30L"),
    ("IndiaAI Mission", "Compute + models + responsible AI", "Subsidised GPUs; Indic eval; privacy kernel", "No fixed amount claimed"),
    ("MeitY DLI / ISM", "Chip design startup, RTL, 28nm edge", "ISA.md + CPU INT8 kernels; FPGA next", "Design-cost support"),
    ("NIDHI-PRAYAS", "Hardware prototype (software-only ineligible)", "SARA Cam enclosure or drone payload running runtime", "Not claimed here"),
])
note("Pick ONE societal lane for first TIDE form. Do not list everything.")

# 5. Immediate Action Plan (next 30 days)
h1("5. 30-Day Action Plan (from 12 Aug 2026)")
table(["Week", "Action", "Owner", "Done?"],
[
    ("Week 1", "Incorporate Pvt Ltd (SPICe+); get CIN", "Rahul", "☐"),
    ("Week 1", "Apply DPIIT Startup India recognition", "Rahul + Founder", "☐"),
    ("Week 1–2", "Pick TIDE CoE (city CCTV or agri drone)", "Founder", "☐"),
    ("Week 2", "Record 2-min demo video (demos_edge.py output)", "Founder", "☐"),
    ("Week 2–3", "Prepare TIDE 2.0 application (use SARA_APPLICATION_DRAFT)", "Founder", "☐"),
    ("Week 3", "Submit TIDE via chosen CoE with demo video + CIN + DPIIT", "Founder", "☐"),
    ("Week 3–4", "Apply IndiaAI Mission compute (compute.indiaai.gov.in)", "Founder", "☐"),
    ("Week 4", "Start APS for Germany MSc (parallel track)", "Founder", "☐"),
])

# 6. Investor/Equity (Amit + Rahul)
h1("6. Equity Commitments (already discussed)")
table(["Person", "Role", "Equity", "Notes"],
[
    ("Amit (8936097968)", "Investor / angel", "3%", "At incorporation, formal agreement at SPICe+"),
    ("Rahul (9304213747)", "Paperwork / incorporation", "3%", "At incorporation, formal agreement at SPICe+"),
    ("Founder (Saurabh)", "Founder", "94%", "Target ≥40% post-Series C"),
])

# 7. Pitch Assets Needed
h1("7. Pitch Assets to Create This Week")
for i in ["2-min demo video (screen-record demos_edge.py + narration)", "1-pager: SARA in 60 seconds (problem → solution → demo → ask)", "TIDE 2.0 application draft (use SARA_APPLICATION_DRAFT.docx)", "Cap table template (founder 94%, Amit 3%, Rahul 3%, ESOP 10% pool)"]:
    bullet("☐ " + i)

# 8. Runbook for Reviewer
h1("8. Reviewer Runbook (copy-paste for grant officer)")
p("```bash")
p("# 1. Clone & setup (2 min)")
p("git clone https://github.com/skmandal3240/SARA && cd SARA")
p("python3 -m venv .venv && source .venv/bin/activate")
p("pip install -r requirements.txt  # CPU torch")
p("")
p("# 2. Run edge tests (must exit 0)")
p("python -m pytest tests/test_edge.py tests/test_privacy.py tests/test_data_catalog.py tests/test_tools.py -q")
p("")
p("# 3. Run edge demos (CCTV see-path, phone agent, mesh)")
p("python demos_edge.py")
p("")
p("# 4. Verify cloud is denied by default")
p("grep -r \"deny\" demos_edge.py  # shows cloud deny logic")
p("```")
p("Expected: 41 tests pass, 3 edge demos pass, no cloud calls without grant.")

doc.save(OUT)
print(f"SAVED {OUT}")