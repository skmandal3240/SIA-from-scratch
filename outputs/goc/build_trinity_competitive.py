#!/usr/bin/env python3
"""Project TRINITY — Competitive Landscape & Moat DOCX (standalone)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_Competitive_Landscape.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Competitive Landscape & Moat\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("August 2026 | SIA / Mandal Holdings | Working Draft\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
p("Framing: TRINITY is not placed against 'another AI startup' — it is the sovereign full-stack alternative. Ratings are indicative [ESTIMATE] unless flagged.")

# ============ 1. COMPETITIVE SET ============
h1("1. Competitive Set")
h2("1.1 Global hyperscalers & frontier labs")
bullet("OpenAI / Anthropic / Google DeepMind: cloud-first, English-first, no sovereign-India commitment. Their edge is the best models.")
bullet("Meta (Llama): strong open weights — but lineage is foreign; sovereignty still an open question.")
h2("1.2 Indian players")
bullet("Government labs (MeitY/C-DAC, IndiaAI consortia): mandate leadership, slower commercial execution.")
bullet("Indian Gen-AI startups: application-layer wrappers — no full-stack from-scratch sovereignty.")
bullet("Chip efforts (Dholera fab + domestic IP): complementary, not competitive — potential partners.")
h2("1.3 The strategic gap")
p("No player today offers all five layers for India: private-on-device, sovereign-owned, from-scratch, Indic-language-first, and government-aligned. That is the white space TRINITY occupies.")

# ============ 2. FEATURE MATRIX ============
h1("2. Capability Matrix (indicative)")
table(["Capability", "Frontier labs", "Meta/OSS", "India labs", "SIA/TRINITY"],
[
    ("On-device inference", "Partial", "Partial", "Limited", "Core"),
    ("Privacy by design / DPDP", "Weak", "Weak", "Partial", "Yes"),
    ("Indic-language focus", "Weak", "Weak", "Partial", "Focus"),
    ("From-scratch stack", "No", "No", "No", "Yes"),
    ("Sovereign/gov-aligned", "Weak", "Neutral", "Yes", "Sovereign"),
    ("Own GPU cluster (long term)", "Yes", "Yes", "Partial", "Plan"),
    ("Edge/offline for masses", "Partial", "Partial", "Limited", "Core"),
    ("Full-stack = layers 1–5", "Partial", "Partial", "No", "Yes"),
])

# ============ 3. MOATS ============
h1("3. Defensible Moats")
h2("3.1 From-scratch IP")
p("The SIA stack — tokenizer, transformer, multimodal encoders, diffusion heads, tools — is original code owned in India. No wrapper fees, no foreign API dependence, patentable ([ESTIMATE] 2–4 filings Y1–2).")
h2("3.2 On-device privacy (DPDP compliance as moat)")
p("Inference at the edge means data never leaves the device. DPDP 2023 forces exactly this posture for Indian enterprises/government; a compliance-by-design is a structural advantage, not a feature.")
h2("3.3 Indian-language & context data")
p("Indian data is the scarce asset; every model needs it, few can build it. SIA Memory + datasets + synthetic pipeline compound over time. This is a data flywheel.")
h2("3.4 Sovereign-government alignment")
p("RFPs for AI infrastructure prefer an Indian owner. TRINITY can win contracts that foreign or wrapper companies cannot — jobs, IP, and security stay in India.")
h2("3.5 Silicon roadmap (DLI/PLI)")
p("Edge NPU aligned to Dholera 28nm and DLI incentive + the nuclear/energy cost curve → structural capex advantage at Scale stage.")

# ============ 4. WEAKNESSES & MITIGATIONS ============
h1("4. Our Gaps vs Competitors (honest)")
table(["Gap", "Competitor strength", "Mitigation"],
[
    ("Model scale vs frontier", "Frontier labs have 100B+ models", "Play edge/on-device niche; cloud tier 70B when warranted"),
    ("Brand & traction", "OpenAI/Google known", "Sovereign positioning + gov contracts win over brand abroad"),
    ("Capital access", "Hyperscalers are cash-rich", "Phase-gated funding + grant pipeline"),
    ("Talent", "Big tech pays top $", "Mission + equity; India talent pool"),
])
p("Net: we do not need to out-compete frontier labs on general chat. We win the underserved 'sovereign + on-device + Indic' segment they structurally cannot serve well.")

# ============ 5. WIN/LOSS FRAMING ============
h1("5. Win/Loss Positioning")
table(["Scenario", "Why TRINITY wins", "Why we might lose", "Mitigation"],
[
    ("Enterprise AI buy", "Privacy, Indic, sovereignty", "Enterprise wants frontier scale", "SIA Pro hybrid (edge+ cloud)"),
    ("Government RFP", "Sovereignty, IP stays, jobs", "Incumbent foreign vendors", "Gov-page + grants pipeline"),
    ("Developer SDK", "Embed-friendly, offline", "Ecosystem gravity of OpenAI", "Open source, tooling, opinions"),
    ("Consumer app", "Private-per device", "f FOMO 'cloud is smarter'", "Edge + cloud fallback"),
])
p("Each loss scenario has a mitigation; the mitigations sum to the product roadmap (hybrid edge/cloud, SDK openness, gov-trust strategy).")

# ============ 6. STRATEGIC TAKEAWAY ============
h1("6. Strategic Takeaway")
p("TRINITY is first-to-market among credible, full-stack, sovereign Indian AI companies. Competition exists at each layer individually but not across all five. The moat is the stack, the data, the DPDP posture, the government trust, and the timeline to national-scale infrastructure. Next proof points: pilot contracts, patent filings, and a deployed edge model with ≥95% tool-call accuracy (SOM-boosters).")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")