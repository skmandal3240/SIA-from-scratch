#!/usr/bin/env python3
"""Project TRINITY — TAM/SAM/SOM Market Sizing DOCX (standalone).
Run: python3 outputs/goc/build_trinity_tam.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_TAM_SAM_SOM.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Market Sizing: TAM / SAM / SOM\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("August 2026 | SIA / Mandal Holdings | Working Draft\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
p("All figures [ESTIMATE] unless flagged [VERIFIED]. Directional for planning; bottom-up vertical sheets follow in the financial workbook.")

# ============ 1. METHODOLOGY ============
h1("1. Methodology")
p("Standard TAM→SAM→SOM funnel. TAM = top-down published India AI spend. SAM = serviceable segments we can actually reach with our delivery model (on-device, sovereign aligned, SDK/agents). SOM = conservative capture (1% of SAM by year 5, phased by pillar). Bottom-up vertical sheets (health/edu/agri/defence/gov) refine SAM by use case in the financial workbook.")

# ============ 2. TAM ============
h1("2. Total Addressable Market (TAM)")
h2("2.1 India AI spend")
table(["Metric", "2026E", "2030E", "Source"],
[
    ("India AI total spend (compute+models+apps+services)", "$12–15B", "$25–40B", "Analyst consensus [VERIFIED direction]"),
    ("Global AI market", "$300–400B", "$800B+", "Multiple analysts [VERIFIED direction]"),
])
p("TAM framing: every rupee India spends on AI — public and private — is addressable by a full-stack Indian player. The current dependency on foreign cloud is the wedge.")
h2("2.2 TAM nuances for TRINITY")
bullet("We are not 'a piece of AI spend' — we can capture multiple layers of it (compute margin + model margin + application margin).")
bullet("Government segment is TAM-relevant and hardest for foreigners — sovereign premium applies.")

# ============ 3. SAM ============
h1("3. Serviceable Addressable Market (SAM)")
h2("3.1 SAM by delivery model")
table(["Segment", "2028E SAM", "Rationale"],
[
    ("Enterprise AI (B2B + gov)", "$3–5B", "SIA Pro + agents + contracts"),
    ("On-device/Indian-language AI", "$1–2B", "SIA Edge; low-end phones offline"),
    ("Developer/SDK embedded AI", "$0.5–1B", "SIA Studio; OEM/IoT"),
    ("TOTAL SAM", "~$4–6B", "Wedge = sovereignty + privacy + Indic"),
])
h2("3.2 What SAM excludes")
bullet("Foreign-cloud-only workloads we can't serve (migrated later)")
bullet("Consumer AI spend in top-tier broadband markets (US/EU)")
bullet("Generic GPU rental markets outside India (deliberate scope)")

# ============ 4. SOM ============
h1("4. Serviceable Obtainable Market (SOM)")
h2("4.1 Capture (5-year)")
table(["Year", "Sam capture", "SOM (illustrative)", "₹ equiv"],
[
    ("Y1", "0%", "$0", "₹0"),
    ("Y2", "0.2%", "$8–12M", "₹70–100 Cr"),
    ("Y3", "0.5%", "$20–30M", "₹170–250 Cr"),
    ("Y5", "1.0%", "$40–60M", "₹330–500 Cr"),
    ("Y10", "2.5%", "$100–150M", "₹830–1,250 Cr"),
])
p("SOM aligns to the funding ladder: ₹100–500 Cr Scale stage maps to year-3-ish SOM. Y5 1% is a conservative, defensible capture rate for a first-mover sovereign player.")
h2("4.2 Bottom-up sanity check (one line)")
p("1% of SAM = ~$50M = roughly 100 enterprise/gov accounts at $0.5M ACV, or 10M consumers at ~$5/yr — internally consistent with the revenue build.")

# ============ 5. VERTICAL BREAKDOWN (future workbook) ============
h1("5. Vertical Breakdown (to be completed in financial workbook)")
table(["Vertical", "TAM element", "SAM driver", "Status"],
[
    ("Enterprise AI", "Large", "PoC pipeline", "Draft"),
    ("Government digital missions", "Large", "Sovereign RFPs", "Draft"),
    ("Healthcare", "Large", "CLINIC + rural", "TODO bottom-up"),
    ("Education", "Large", "Vernacular", "TODO bottom-up"),
    ("Agriculture", "Large", "Advisory on-device", "TODO bottom-up"),
    ("Defence/security", "Large", "Sovereign edge", "TODO bottom-up"),
])
p("Bottom-up sheets per vertical with device counts, price points, and adoption assumptions.")

# ============ 6. SOURCES ============
h1("6. Sources & Assumptions")
bullet("IndiaAI Mission: 34,000+ GPUs, 367+ datasets, 20 sovereign models [VERIFIED Aug 2026]")
bullet("India AI market >25%/yr growth [VERIFIED direction — analysts]")
bullet("DPDP 2023 as compliance driver [VERIFIED regulation]")
bullet("Access: 900M+ users rely on lower-bandwidth/mobile-first internet [VERIFIED context]")
p("Each number above is an estimate for planning; validate with named reports when presenting externally.")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")