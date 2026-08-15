#!/usr/bin/env python3
"""SARA — TIDE 2.0 Grant Application Draft (edge AI, CCTV/drone lane)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "SARA_TIDE2_Application_Draft.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
s = doc.styles["Normal"]
s.font.name = "Calibri"
s.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = NAVY; r.font.size = Pt(16)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = GOLD; r.font.size = Pt(13)

def p(t): doc.add_paragraph(t)
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def note(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.color.rgb = GREY; r.font.size = Pt(9)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hh in enumerate(headers): t.rows[0].cells[i].text = hh
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row): t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()

# Cover
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("SARA — TIDE 2.0 GRANT APPLICATION DRAFT\n"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
r = doc.add_paragraph("MeitY Technology Incubation & Development of Entrepreneurs\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(13); r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Applicant: Saurabh Mandal · SARA AI (proposed Pvt Ltd) · Virar, Maharashtra · 12 August 2026\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(10); r.runs[0].font.color.rgb = GREY
note("Repo: https://github.com/skmandal3240/SARA — all code in-house, Apache-2.0. Demo runs on CPU, no GPU, no cloud keys.")

# 1. Executive Summary
h1("1. Executive Summary (100 words)")
p("SARA is an India-first, from-scratch multimodal AI that runs on the device — not in the cloud. Its transformer (text, vision, audio, code, tools, agents) is built from scratch in this repo (Apache-2.0). SARA's edge runtime profiles (CCTV, drone, phone, laptop) execute inference locally; mesh offloads only when local+peer capacity is saturated; cloud is denied by default unless the user grants consent. We seek TIDE 2.0 prototype grant to build two device prototypes (SARA Cam + SARA Drone payload) running the same runtime, and to train beyond-nano weights on IndiaAI GPUs for Indic languages and edge vision.")

# 2. Problem
h1("2. Problem (TIDE Societal Lane: Infrastructure/Transport + Agriculture)")
bullet("900M+ Indians lack reliable high-speed internet; cloud AI fails offline.")
bullet("DPDP 2023: sending camera/mic data to foreign servers is a compliance risk.")
bullet("City CCTV and farm drones today record → upload → cloud AI → latency/breach.")
bullet("No Indian edge AI runtime exists that is from-scratch, multimodal, agentic, and privacy-by-design.")

# 3. Solution
h1("3. Solution — SARA Edge Runtime")
bullet("Same binary, different profile YAML: CCTV (see/alert), Drone (see/radio), Phone (talk/code/vault).")
bullet("On-device inference: INT8 quant + layer paging for 2–4 GB RAM class; no model leaves device.")
bullet("Mesh: task-DAG scheduler places work on cheapest capable peer; payloads = embeddings/tokens, not raw video.")
bullet("Privacy kernel: grants/vault/audit/local-learn — raw data never leaves without user preview+approve.")
bullet("Agents: sandboxed Python exec, ReAct loop, swarm delegation — core runtime, not API wrapper.")

# 4. Innovation
h1("4. Innovation & Technical Differentiators")
table(["Differentiator", "How SARA Does It"],
[
    ("From-scratch multimodal transformer", "10.8M params, GQA, RoPE, RMSNorm, SwiGLU, cross-attn vision/audio"),
    ("Unified agent + tool protocol", "`<|tool_call|>` JSON, sandboxed Python, retry on error, swarm delegation"),
    ("Edge-first profiles", "YAML caps: RAM, TOPS, modalities, cloud-policy; same runtime, zero #ifdef"),
    ("Mesh offload (not cloud)", "DAG scheduler → peer devices; embeddings only; cloud denied by default"),
    ("Privacy kernel in code", "Grants (capability preview), Vault (encrypted memory), Audit (signed log), Local learn (LoRA)"),
    ("IndiaAI-ready training", "Dataset adapters (IndicCorp, IndicVoices, AI4Bharat); HF load, no vendored TB"),
])

# 5. Societal Impact (TIDE lane)
h1("5. Societal Impact (Lane: Infrastructure/Transport — CCTV; Agriculture — Drone)")
table(["Lane", "Device", "What It Does Locally", "Impact"],
[
    ("CCTV / City", "SARA Cam", "See (detect/caption), alert on anomaly, 24/7 INT8, mesh to hub", "No raw video leaves premises; privacy-compliant; low bandwidth"),
    ("Agri / Drone", "SARA Drone payload", "See crop health, short talk alert, radio mesh to farmer phone", "Offline operation in fields; farmer data stays on device"),
])
note("Pick ONE lane for the form. CCTV is faster to prototype (no flight certification). Drone adds agri lane for Phase B.")

# 6. Prototype Plan (TIDE deliverable)
h1("6. Prototype Deliverables (18 months)")
table(["Milestone", "Months", "What We Ship"],
[
    ("M1: Incorporation + DPIIT + CoE pick", "0–1", "CIN, DPIIT number, TIDE CoE letter of intent"),
    ("M2: SARA Cam prototype (box + runtime)", "2–4", "CCTV box running see-path; cloud denied demo video"),
    ("M3: SARA Drone payload prototype", "4–8", "Drone compute module running see+radio mesh; field test"),
    ("M4: IndiaAI training run", "6–12", "Indic + code + vision fine-tune on subsidised GPUs"),
    ("M5: Mesh in a shop/house (3+ devices)", "10–14", "CCTV + phone + hub mesh; task DAG placement log"),
    ("M6: TIDE prototype grant closeout", "18", "Two devices, runtime, trained weights, audit report"),
])

# 7. Use of Funds
h1("7. Use of Funds — Ask: up to ₹30,00,000 (TIDE prototype grant)")
table(["Item", "Amount", "Share"],
[
    ("IndiaAI GPU training credits (beyond nano)", "₹10,00,000", "33%"),
    ("SARA Cam hardware (SoC + enclosure + sensors)", "₹7,00,000", "23%"),
    ("SARA Drone payload hardware", "₹5,00,000", "17%"),
    ("Team stipend (founder + 1 engineer)", "₹5,00,000", "17%"),
    ("Legal / DPIIT / compliance", "₹2,00,000", "7%"),
    ("Travel / field tests / misc", "₹1,00,000", "3%"),
])

# 8. Team
h1("8. Team")
table(["Role", "Name", "Background"],
[
    ("Founder / AI Architect", "Saurabh Mandal", "BE CSE 2025, CGPA 7.22; built SARA from-scratch (10.8M params, 9/9 demos), ASTRO, ALICE, SHADE"),
    ("Paperwork / Incorporation", "Rahul (9304213747)", "Equity 3% at incorporation; SPICe+, DPIIT, compliance"),
    ("Investor / Angel", "Amit (8936097968)", "Equity 3% at incorporation; seed capital bridge"),
])
note("Advisors: to be assigned by TIDE CoE post-selection.")

# 9. Pre-Submission Checklist
h1("9. Pre-Submission Checklist (as of 12 Aug 2026)")
for i in [
    "☐ Private Limited incorporated (SPICe+); CIN obtained",
    "☐ DPIIT Startup India recognition applied (certificate pending)",
    "☐ ≥51% Indian ownership confirmed in cap table",
    "☐ TIDE CoE picked (city CCTV lane recommended for speed)",
    "☐ CoE letter of intent obtained",
    "☐ 2-min demo video recorded (demos_edge.py output + narration)",
    "☐ Application form drafted (this document)",
    "☐ Founder APS application started (parallel Germany MSc track)",
]:
    bullet(i)

# 10. Reviewer Runbook (copy-paste)
h1("10. Reviewer Runbook (CPU, no GPU, no cloud keys)")
p("```bash")
p("git clone https://github.com/skmandal3240/SARA && cd SARA")
p("python3 -m venv .venv && source .venv/bin/activate")
p("pip install -r requirements.txt")
p("python -m pytest tests/test_edge.py tests/test_privacy.py tests/test_data_catalog.py tests/test_tools.py -q")
p("python demos_edge.py")
p("```")
p("Expected: 41 tests pass, 3 edge demos pass (CCTV see-path cloud denied, Phone agent cloud denied, 2-node mesh DAG). Nano captions weakly trained — expected. Review: **does work stay on chip, cloud stay off?**")

doc.save(OUT)
print(f"SAVED {OUT}")