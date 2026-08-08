#!/usr/bin/env python3
"""Project TRINITY — FULLY DETAILED 30-page master investor report.
Chunked builder: this file grows by appending section files (part2, part3...).
Run: python3 outputs/goc/build_trinity_master.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
def add_figure(path, caption=None):
    from docx.shared import Inches
    if Path(path).exists():
        if caption:
            doc.add_paragraph(caption)
        doc.add_picture(str(path), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "Project_TRINITY_Master_Report.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = NAVY
r = doc.add_paragraph("India's Sovereign AI Infrastructure Company — Master Investor Report\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("August 2026 | SIA / Mandal Holdings | Working Draft — v2 (Appendix A/B/C merged)\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
p("Legend: [VERIFIED] = sourced/confirmed. [ESTIMATE] = illustrative planning figure. [TODO] = requires founder input or professional advice.")

# ============ 1. EXEC SUMMARY ============
h1("1. Executive Summary")
p("Project TRINITY proposes building a complete AI technology stack for India across five strategic layers: Compute Infrastructure, Data Infrastructure, Foundation Models, AI Agents & AI Operating System, and AI Applications. The long-term vision is to create AI infrastructure and products serving enterprise, government, developers, and consumers — positioned not as a typical AI startup, but as India's sovereign AI infrastructure company.")
p("The delivery vehicle is SIA — a family of on-device and cloud AI models (Nano 0.5–1B to Cloud 70B+) sharing one tokenizer, one architecture, one tool-calling interface, one memory format, and one runtime. SIA is an AI operating system: private by design, built from scratch in India, embedded in every tier of hardware from wearables to data centers.")
p("Phased funding: MVP ₹2–10 Cr → Growth ₹25–100 Cr → Scale ₹100–500 Cr → National ₹500–2,000 Cr → Global ₹2,000 Cr+. Near-term ask: ₹2–10 Cr to deliver a working flagship demo, five enterprise/government pilots, and India's first credible on-device sovereign AI stack.")# ============ 2. VISION / MISSION / WHY NOW ============
h1("2. Vision, Mission & Why Now")
h2("Vision")
p("Build AI infrastructure, models, and applications that enable India to strengthen domestic AI capabilities while competing globally.")
h2("Mission")
p("Develop an integrated AI ecosystem covering infrastructure, models, developer platforms, and applications — the complete stack required for India to become an AI sovereign nation.")
h2("Why Now")
bullet("IndiaAI Mission: 34,000+ GPUs, 367+ datasets, 20 sovereign models funded [VERIFIED Aug 2026]")
bullet("DPDP 2023 makes on-device privacy a regulatory advantage")
bullet("Jevons paradox: cheaper AI drives exploding usage — cost leaders win")
bullet("Global race (US, China, EU) has no Indian full-stack player at sovereign scale")

# ============ 3. INDIA'S AI OPPORTUNITY ============
h1("3. India's AI Opportunity")
h2("3.1 Market context [VERIFIED + ESTIMATE]")
bullet("IndiaAI compute exceeds 34,000 GPUs; AI Kosh 367+ datasets; 500+ foundation-model proposals; 20 sovereign models to be supported [VERIFIED, DD News/CXO Today Aug 2026]")
bullet("India AI market growing >25%/yr; on-device/edge AI the fastest segment [ESTIMATE]")
bullet("900M+ lack reliable connectivity — on-device AI is the only deliverable path")
h2("3.2 The five-pillar prize")
p("Each pillar creates its own margin layer: chips, clouds, models, agents, and applications. Owning them together is the sovereign moat; each strengthens the others (five-pillar flywheel).")
table(["Pillar", "What India gets", "TRINITY lever"],
[
    ("1. Compute", "AI sovereignty in silicon + cloud", "Rent → own cluster timeline"),
    ("2. Data", "Indian-language IP", "SIA Memory + datasets"),
    ("3. Models", "Domestic foundation models", "SIA family (Nano→Cloud)"),
    ("4. Agents & OS", "National AI runtime", "AI OS + tool interface"),
    ("5. Applications", "Jobs + public value", "Health/edu/agri/defence"),
])

# ============ 4. AI'S FIVE PILLARS ============
h1("4. AI's Five Pillars Framework")
table(["Pillar", "Scope"],
[
    ("1. Compute", "GPU cloud, AI clusters, data centers, chips (CPU/GPU/TPU/NPU/photonic)"),
    ("2. Data", "Indian datasets, synthetic data, enterprise/government datasets, marketplace"),
    ("3. Models", "SIA family: Nano 0.5–1B / Edge 2–4B / Pro 8–14B / Cloud 70B+; vision, speech, multimodal"),
    ("4. Agents & OS", "AI OS, autonomous agents, coding agents, enterprise agents, robotics agents"),
    ("5. Applications", "Healthcare, education, agriculture, manufacturing, defence, government, consumer AI"),
])

# ============ 5. PILLAR 1 — COMPUTE ============
h1("5. Pillar 1 — Compute Infrastructure")
p("Infrastructure follows a staged build-vs-rent strategy:")
table(["Horizon", "Action"],
[
    ("Years 1–2", "Rent GPU (Yotta, E2E, international) — focus on model development"),
    ("Years 2–4", "Build software infrastructure: distributed training, inference platform, MLOps, data pipelines"),
    ("Years 4–6", "Small private GPU cluster (~100–500 GPUs) for predictable workloads"),
    ("Years 6–10", "Hyperscale AI campus: thousands of GPUs, high-speed networking, dedicated substations, liquid cooling, peta-scale storage"),
])
h2("Chip strategy")
bullet("CPU for OS; GPU for training today; TPU-class tensor throughput")
bullet("NPU for local inference (phones, laptops, robots, cars, IoT)")
bullet("Optical/photonic for future data-center LLM workloads")
bullet("Edge NPU roadmap aligns with DLI/PLI + Dholera 28nm [VERIFIED context]")

# ============ 6. PILLAR 2 — DATA ============
h1("6. Pillar 2 — Data Infrastructure")
bullet("Indian-language + Indian-context datasets (the untapped moat)")
bullet("SIA Memory: consistent, personalized cross-device experience; compatible memory formats across tiers")
bullet("Synthetic data pipeline to scale scarce domain datasets")
bullet("Public + enterprise + government datasets; future data marketplace")# ============ 7. PILLAR 3 — MODELS ============
add_figure("outputs/goc/diagrams/02_sia_model_family.png", "Figure 2: SIA model family — one AI OS, four hardware tiers, one shared stack.")
h1("7. Pillar 3 — Foundation Models (SIA)")
p("SIA is an AI operating system — a family of models, not one giant model, all sharing the same stack:")
table(["Tier", "Size", "Target"],
[
    ("SIA Nano", "0.5B – 1B", "Wearables, IoT, microcontrollers"),
    ("SIA Edge", "2B – 4B", "Phones, PCs, laptops"),
    ("SIA Pro", "8B – 14B", "Workstations, local GPUs"),
    ("SIA Cloud", "70B+", "Data centers — advanced reasoning"),
])
p("Shared across all tiers: tokenizer, architecture, tool-calling interface, memory formats, and a common runtime that synchronizes knowledge. From-scratch nano transformer already trained on CPU (loss 2.78, 9/9 demos); LoRA fine-tune pipeline ready (Gemma 3 1B base) [VERIFIED].")

# ============ 8. PILLAR 4 — AGENTS & OS ============
h1("8. Pillar 4 — AI Agents & Operating System")
bullet("SIA Agent: autonomous task execution")
bullet("SIA Studio / SDK: developers embed SIA into apps and IoT products")
bullet("Common tool-calling interface (calc/now/file tools via [[tool:name(args)]]) [VERIFIED]")
bullet("Coding agents (Claude Code / Codex class), enterprise agents, workflow automation")

# ============ 9. PILLAR 5 — APPLICATIONS ============
h1("9. Pillar 5 — Applications & Ecosystem")
table(["Module", "Role"],
[
    ("SIA Edge", "Local on-device AI model (delivery vehicle)"),
    ("SIA Core", "Central orchestration and synchronization"),
    ("SIA Cloud", "Advanced reasoning when local isn't enough"),
    ("SIA Studio / SDK", "Developer embed — apps and IoT"),
    ("SIA Memory", "Consistent personalized experience across devices"),
    ("SIA Vision", "Multimodal perception, recognition, synthesis"),
    ("SIA Voice", "Speech recognition + synthesis"),
    ("SIA Agent", "Autonomous task execution"),
])
p("Target verticals: healthcare, education, agriculture, manufacturing, defence, space, government, consumer AI.")

# ============ 10. TECHNOLOGY ARCHITECTURE ============
add_figure("outputs/goc/diagrams/01_5layer_ai_stack.png", "Figure 1: The 5-Layer AI stack — coding agents to chip + energy.")
add_figure("outputs/goc/diagrams/03_memory-agents-loop.png", "Figure 3: SIA memory / agent / tool loop with DPDP privacy boundary.")
h1("10. Technology Architecture")
p("The 5-Layer AI stack:")
table(["Layer", "Slice"],
[
    ("L1", "Claude Code / Codex — coding agents"),
    ("L2", "Applications — All"),
    ("L3", "Models — foundation + SIA family"),
    ("L4", "Infrastructure — data centers"),
    ("L5", "Chip + Energy (nuclear path)"),
])
p("Outcomes: SIA Everywhere (models embedded in every device) and Root India / Build India (IP, jobs, chips, energy in India).")
h2("Verification from the live repo")
bullet("9/9 demo gauntlet passing: text, code, vision, audio-listen, audio-gen, image (VAE-decoded), video, tools, quantize [VERIFIED]")
bullet("Audio→language loop live (hear → describe → reply → speak back) [VERIFIED]")

# ============ 11. PRODUCT ROADMAP ============
h1("11. Product Roadmap (2026–2035)")
table(["Phase", "Objective", "Horizon"],
[
    ("1", "Prototype & MVP (SIA Edge on-device companion)", "2026–2027"),
    ("2", "Foundation models (Nano/Edge trained, Pro LoRA)", "2027–2029"),
    ("3", "Enterprise platform (agents, SDK, memory)", "2029–2031"),
    ("4", "National AI infrastructure (cluster → campus)", "2031–2035"),
    ("5", "Global expansion", "2035+"),
])

# ============ 12. COMPETITIVE LANDSCAPE ============
h1("12. Competitive Landscape")
table(["Capability", "OpenAI", "Google", "India labs", "SIA/TRINITY"],
[
    ("On-device inference", "Partial", "Partial", "-", "Core"),
    ("Privacy by design/DPDP", "-", "-", "-", "Yes"),
    ("Indic languages", "Weak", "Weak", "Partial", "Focus"),
    ("From-scratch stack", "No", "No", "No", "Yes"),
    ("Sovereign/gov contracts", "Hard", "Hard", "Yes", "Sovereign pitch"),
    ("Own GPU cluster (long term)", "Yes", "Yes", "Partial", "Plan"),
])
p("Moat: from-scratch IP, on-device privacy, Indian languages/context, silicon roadmap (DLI/PLI), sovereign-government alignment.")

# ============ 13. WHY INDIA HAS AN ADVANTAGE ============
h1("13. Why India Has an Advantage")
bullet("Engineering talent density + low cost of skilled labour")
bullet("Data sovereignty mandate (DPDP) creates a compliant-by-design moat")
bullet("Government demand: IndiaAI, Digital India, defence, state digital missions")
bullet("Scale: 1.4B people, 900M+ offline — on-device is the only path")
bullet("Semiconductor runway: Dholera 28nm, DLI/PLI, RISC-V domestic IP")

# ============ 14. MARKET OPPORTUNITY (TAM/SAM/SOM) ============
h1("14. Market Opportunity (TAM/SAM/SOM)")
h2("TAM [ESTIMATE]")
p("India AI total spend 2026E ~$12–15B (policy + industry reports); global AI market $800B+ by 2030 (multiple analysts).")
h2("SAM [ESTIMATE]")
p("Serviceable = enterprise AI + on-device Indian-language AI + developer/SDK: ~$4–6B by 2028.")
h2("SOM [ESTIMATE]")
p("Conservative 1% of SAM by year 5 ≈ $40–60M ARR (₹330–500 cr), phased by pillar. Aligns with the ₹100–500 cr funding stage.")
p("Note: directional; vertical bottom-up sheets (health/edu/agri/defence) to follow in the financial workbook.")# ============ 15. BUSINESS MODEL ============
h1("15. Business Model")
bullet("AI API usage (SIA Cloud tier)")
bullet("Enterprise subscriptions (SIA Pro + agents)")
bullet("GPU cloud (post-Year 4 private cluster)")
bullet("Managed AI services + licensing")
bullet("Professional services / SDK embed fees")
h2("Illustrative revenue build [ESTIMATE]")
table(["Line", "Y1", "Y3", "Y5", "Y10"],
[
    ("SIA Edge consumer", "₹0", "₹2Cr", "₹8Cr", "₹50Cr"),
    ("SIA API", "₹0", "₹1Cr", "₹10Cr", "₹60Cr"),
    ("Enterprise subscriptions", "₹0", "₹5Cr", "₹25Cr", "₹150Cr"),
    ("Government contracts", "₹0", "₹3Cr", "₹15Cr", "₹80Cr"),
    ("GPU cloud + services", "₹0", "₹0", "₹10Cr", "₹120Cr"),
    ("TOTAL (illustrative)", "₹0", "~₹11Cr", "~₹68Cr", "~₹460Cr"),
])

# ============ 16. GTM ============
h1("16. Go-to-Market Strategy")
table(["Seq", "Channel", "Mechanic", "KPI"],
[
    ("1", "Developers (SIA Studio/SDK)", "Open-source + SDK", "#Downloads"),
    ("2", "Startups (India)", "Freemium + API", "Signups, API $"),
    ("3", "SMEs", "Managed pilots", "Pilots → paid"),
    ("4", "Enterprises", "Direct PoC", "PoC → contract"),
    ("5", "Government", "Sovereign RFPs, grants", "Contract #"),
])

# ============ 17. FINANCIALS ============
h1("17. Financial Model")
h2("Unit economics (indicative) [ESTIMATE]")
table(["Metric", "Target", "Notes"],
[
    ("Gross margin", "70–85%", "Edge ~zero marginal cost; cloud GPU cost"),
    ("Enterprise CAC", "₹2–5 lakh", "Direct sales + gov partnerships"),
    ("Enterprise LTV", "₹30+ lakh / 4 yrs", "Multi-year renewals + agents"),
    ("LTV/CAC", ">3 (target 4–6×)", "Healthy SaaS"),
    ("Churn", "<5%/yr enterprise", "Switch costs high"),
])
p("3-statement model (P&L/BS/CF) to be built in the financial workbook; the revenue table in §15 is the shape.")

# ============ 18. FUNDING STRATEGY ============
h1("18. Funding Strategy")
table(["Stage", "Goal", "Illustrative Range"],
[
    ("MVP", "Prototype", "₹2–10 Cr"),
    ("Growth", "Foundation models", "₹25–100 Cr"),
    ("Scale", "AI cloud", "₹100–500 Cr"),
    ("National", "Infrastructure", "₹500–2,000 Cr"),
    ("Global", "Expansion", "₹2,000 Cr+"),
])
p("Illustrative planning ranges only — actual funding depends on milestones, eligibility, investor appetite, and program requirements. Company registered in Bengaluru, Karnataka.")
h2("Government sources")
p("IndiaAI Mission, Startup India, MeitY (TIDE 2.0 / SAMRIDH), ANRF, DST (SISFS, NIDHI-PRAYAS), Digital India, Semiconductor Mission (DLI/PLI), IN-SPACe, state startup missions (incl. Karnataka state schemes).")
h2("Private sources")
p("Angels, deep-tech VCs, AI-focused VCs, strategic investors, hyperscalers, semiconductor companies, sovereign wealth funds.")

# ============ 19. RISKS ============
h1("19. Risk Analysis")
table(["Risk", "Mitigation"],
[
    ("Capital intensity", "Staged rent→build; milestone-gated funding"),
    ("Competition (hyperscalers)", "Sovereignty + on-device privacy differentiation"),
    ("Technology execution", "Working from-scratch stack; LoRA pipeline"),
    ("Regulatory changes", "DPDP alignment is an advantage"),
    ("Talent acquisition", "Indian dev pool + remote-first"),
])

# ============ 20. AGI POSITION ============
h1("20. AGI Roadmap & Scientific Position")
p("Evidence-based framing only. Current AI capabilities are narrow; scientific limitations remain in reasoning, memory, compute, data, and robotics. Project TRINITY is designed to build the infrastructure, models, and ecosystem that position the company to contribute to increasingly capable AI systems over time. Whether AGI is achieved remains scientifically uncertain and cannot be guaranteed.")

# ============ 21. TEAM ============
h1("21. Team & Hiring Plan")
bullet("Founder: Saurabh Mandal — built the from-scratch SIA framework (tokenizer, transformer, multimodal, tools), 3 prior products (ASTRO, ALICE, SIA)")
bullet("Legal entity: Indian Private Limited Company via SPICe+, registered office Bengaluru, Karnataka [VERIFIED]")
bullet("Hiring: Y1: 1–2 engineers + 1 product/design; Y2: 3–5 (ML, infra); Y3+: platform + enterprise")
table(["Year", "Roles", "Budget (Indicative)"],
[
    ("1", "2 eng + 1 product/design", "₹25–40 lakh/yr"),
    ("2", "ML + infra + 1 BD", "₹60–80 lakh/yr"),
    ("3", "Platform + enterprise", "₹1.2–1.5 cr/yr"),
    ("4+", "Full team; ops + cluster engineers", "₹3 cr/yr+"),
])
p("Advisory candidates [TODO]: ex-IndiaAI/MeitY leader; ex-cloud exec; NLP/edge academic.")

# ============ 22. VALUATION ============
h1("22. Company Valuation Strategy")
p("Valuation not revenue-multiple alone; blend:")
bullet("DCF on long horizon with phase-gated milestone discounting")
bullet("Comparables: AI-infra (CoreWeave, Groq), applied AI (Palantir-class), Indian AI startups")
bullet("Precedent transactions + VC method with dilution waterfall")
p("Leveraging inputs: signed gov contracts, GPUs deployed, model quality evals (≥95% tool-accurate), patents granted. [TODO] full model once DPIIT + first pilot are real.")

# ============ 23. APPENDIX A — INVESTOR DD (KEY Q&A) ============
h1("23. Investor Due Diligence (Key Q&A)")
q = [
    ("Why now?", "IndiaAI Mission funding 20 sovereign models; DPDP privacy advantage; Jevons paradox favors cost leaders; no Indian full-stack sovereign player."),
    ("Why this team?", "Founder built the entire from-scratch stack alone — tokenizer, transformer, multimodal, tools — plus 3 prior products and a complete grant pipeline."),
    ("Why this market?", "India AI growing >25%/yr; 900M+ users offline; government demand for sovereign AI; no credible domestic competitor."),
    ("What is the competitive advantage?", "From-scratch IP, on-device privacy, Indian-language data, silicon roadmap, sovereign-government alignment."),
    ("How will capital be deployed?", "MVP ask ₹2–10 Cr: ~35% compute, ~30% team, ~20% data/fine-tuning, ~15% legal/SDK."),
    ("What milestones unlock the next round?", "MVP demo + 5 pilots + 1 government contract → Growth round (₹25–100 Cr) for foundation models."),
]
for question, answer in q:
    h2(question)
    p(answer)

# ============ 24. INVESTMENT ASK ============
h1("24. Investment Ask & Use of Funds")
p("Near-term ask: ₹2–10 Cr MVP phase — team, prototype, compute.")
table(["Use", "Share"],
[
    ("Team (founder + 1–2 engineers)", "~30%"),
    ("GPU compute (rented: Yotta/E2E/Colab)", "~35%"),
    ("Data + fine-tuning (Indian languages)", "~20%"),
    ("Legal, compliance, filings, SDK tooling", "~15%"),
])

# ============ 25. CLOSING ============
h1("25. Closing Vision & Call to Action")
p("We are building the complete AI technology stack required for India to become an AI sovereign nation — SIA, an AI operating system embedded in every device, with India as its root. Join us in building the infrastructure, models, and applications that keep India's data in India, India's intelligence in India, and India's AI future in Indian hands.")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")