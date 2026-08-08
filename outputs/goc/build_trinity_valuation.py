#!/usr/bin/env python3
"""Project TRINITY — Valuation Strategy Outline DOCX (standalone)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_Valuation_Outline.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Valuation Strategy — Outline\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("August 2026 | SIA / Mandal Holdings | Working Draft\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
p("Valuation approach, methods, levers, and a cap-table path. Figures [ESTIMATE] unless flagged. Final model to be built once DPIIT registration and first pilot exist.")

# ============ 1. VALUATION PHILOSOPHY ============
h1("1. Valuation Philosophy")
bullet("Not a revenue-multiple on year-1 revenue — there is none. This is infrastructure + frontier-tech: value by milestone-backed assets and strategic optionality.")
bullet("Sovereign/strategic premium: Indian ownership of AI stack is an asset class of its own for government and strategic investors.")
bullet("Phased re-rate: each milestone (pilot, patent, model quality, GPU cluster) justifies a step up in valuation.")

# ============ 2. METHODS (blend) ============
h1("2. Valuation Methods (blend)")
table(["Method", "How", "Place in blend"],
[
    ("DCF", "Long-horizon 10-yr cash flows from financial model, phase-gated discounting", "Anchor for later stages"),
    ("Comparables", "AI-infra (CoreWeave, Groq), applied AI (Palantir-class), Indian AI startups", "Cross-check"),
    ("Precedent transactions", "Recent India AI deals / acquisition prices in segment", "Sentiment check"),
    ("VC / risk-adjusted method", "Exit value → with dilution waterfall and required return", "Round-size shaping"),
])
p("No single method is decisive; triangulation is the honest answer.")

# ============ 3. VALUATION LEVERS (what changes the number) ============
h1("3. Valuation Levers")
table(["Lever", "Impact", "Trigger"],
[
    ("Signed gov/enterprise contracts", "High — removes adoption risk", "1–2 contracts"),
    ("Deployed edge model (≥95% tool-acc)", "High — technical proof", "Model eval pass"),
    ("GPU cluster in operation", "High — asset-backed", "100–500 GPU"),
    ("Patents granted", "Medium", "First 2–4 filings"),
    ("DPIIT/Startup India + grants", "Medium", "Certificates"),
    ("Revenue/ARR", "High after its start", "First paid pilot"),
])
p("Every lever is tied to a milestone we control.")

# ============ 4. ILLUSTRATIVE STAGE VALUATIONS ============
h1("4. Illustrative Stage Ranges [ESTIMATE]")
table(["Stage", "Ask", "Illustrative post-money", "Multiple basis"],
      [
          ("Pre-seed/MVP", "₹2–10 Cr", "₹15–40 Cr", "Milestone/asset-based"),
          ("Seed/Growth", "₹25–100 Cr", "₹100–250 Cr", "Pilots + 1 gov contract"),
          ("Series A / Scale", "₹100–500 Cr", "₹500–1,500 Cr", "ARR + infra asset"),
          ("Series B / National", "₹500–2,000 Cr", "₹2,000–6,000 Cr", "National infra + revenue"),
          ("Global", "₹2,000 Cr+", "TBD", "Expansion multiple"),
      ])
p("Ranges intentionally wide: they will tighten as levers (§3) fire. The ask follows the model, not the reverse.")

# ============ 5. CAP TABLE & DILUTION PATH ============
h1("5. Cap Table & Dilution Path (illustrative)")
bullet("Incorporation: 100% founder."),
bullet("Seed round target: sell 15–20% (ESOP 10–15% carved alongside).")
bullet("Series A target: 15–20% dilution.")
bullet("Series B: 12–15%. Series C+: 10% (discipline — retain control).")
bullet("Founder target after Series C: ≥40% (with ESOP, bridge, warrants).")
p("Full dilution waterfall computed when cap table is live (post-incorporation).")

# ============ 6. TERM-SHEET CONSIDERATIONS ============
h1("6. Term-Sheet Considerations")
bullet("India Pvt Ltd — typical: pari-passu, liquidation preference (1x non-participating okay), pro-rata rights.")
bullet("Protect founder: reasonable drag-along, avoid participating preferred, avoid unvested cliff pain.")
bullet("Board: founder + 1 investor + 1 independent at early rounds; later add government/strat directors.")
p("[TODO] legal advice before signing any term sheet.")

# ============ 7. EXIT SCENARIOS ============
h1("7. Exit Scenarios [ESTIMATE]")
table(["Path", "How", "Scale"],
      [("Strategic acquisition", "OEM/hyperscaler/government entity buys the stack", "$500M–2B+"),
       ("IPO (India/global)", "Public markets post-scale", "$2B+"),
       ("Long-term sovereign entity", "National AI infrastructure owner (NMP-style)", "Infra value")])
p("No near-term exit intent; the document just shows the lane is open.")

# ============ 8. WHAT'S SKIPPED / NEXT ============
h1("8. What's Skipped (YAGNI) & Next Steps")
bullet("No definitive valuation yet — real model once DPIIT + pilots.")
bullet("No equity split agreed with advisors yet — recommend advisor grant 0.5–1% each (vested).")
p("Next: after incorporation — build live cap table, valuations scenario workbook (xlsx), and aligned term-sheet review.")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")