#!/usr/bin/env python3
"""TIDE 2.0 application draft DOCX — fill-ready for SIA/TRINITY."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TIDE2_Application_Draft_SIA.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
s = doc.styles["Normal"]
s.font.name = "Calibri"
s.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY; r.font.size = Pt(16)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD; r.font.size = Pt(13)

def p(t): doc.add_paragraph(t)
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def note(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.color.rgb = GREY; r.font.size = Pt(9)

# Cover
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("TIDE 2.0 GRANT APPLICATION — DRAFT\n"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY
r = doc.add_paragraph("MeitY Technology Incubation & Development of Entrepreneurs\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(14); r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Company: TRINITY / SIA AI · Bengaluru, Karnataka · Founder: Saurabh Mandal\n"); r.alignment = WD_ALIGN_PARAGRAPH.CENTER; r.runs[0].font.size = Pt(10); r.runs[0].font.color.rgb = GREY
note("[FILL-IN]: amounts in ₹, dates, incubator partner (51 TIDE incubators — nearest: Bengaluru). Deadlines rolling — check meity.gov.in + Startup India portal.")

# 1. Summary
h1("1. 100-Word Executive Summary (paste-ready)")
p("TRINITY is an India-first AI company building a private, on-device AI companion (SIA). Unlike cloud chatbots that send user data to foreign servers, SIA's entire stack — tokenizer, transformer, vision, audio, and code models — is built from scratch in India and runs on consumer devices (phone, laptop, edge hardware). The product is a vernacular AI assistant that works offline, protecting privacy and enabling Indian-language access. We seek TIDE 2.0 incubation support to (1) fine-tune Indian-language small models, (2) ship our first revenue app, and (3) move to Indian edge silicon.")
note("Word count: ~95 — adjust if portal caps differ.")

# 2. Problem
h1("2. Problem")
bullet("900M+ Indians lack reliable high-speed internet; cloud AI is unusable offline.")
bullet("DPDP 2023 compliance: data leaving devices is a regulatory risk — on-device is the answer.")
bullet("Existing assistants are English-first, cloud-locked, and not built for Indian languages/contexts.")

# 3. Solution
h1("3. Solution")
bullet("SIA Core: from-scratch multimodal transformer (text, audio listen/gen, vision, video, code) — all code in repo, no black boxes.")
bullet("SIA Edge: INT4/INT8 quantized small models runnable on 8GB devices (Gemma 3 1B LoRA route ready).")
bullet("SIA Apps: vernacular AI support platform as first revenue engine — one profitable app funds R&D (Pillar 1 rule).")

# 4. Innovation
h1("4. Innovation & Differentiators")
for i in ["From-scratch stack (no OpenAI/Google wrappers) — full IP ownership, no API fees.",
          "On-device privacy: inference + data stay local (vs cloud chatbots).",
          "Audio-native: listens to music/voice, generates audio — multimodal by design.",
          "India-first: Hindi/Indian-language tokenizer, Indian context, Indian silicon roadmap."]:
    bullet(i)

# 5. Market
h1("5. Market")
bullet("India AI market growing >25%/yr; edge/on-device AI is the fastest segment.")
bullet("B2C: freemium privacy-first companion. B2B: white-label edge AI for OEMs.")
bullet("No credible Indian on-device multimodal assistant today [VERIFIED gap].")

# 6. Business Model
h1("6. Business Model")
bullet("Freemium app (privacy = premium). B2B licensing. Grants as non-dilutive fuel.")
bullet("Rule: 30–50% margin reinvested into models and compute.")

# 7. Team
h1("7. Team")
p("Founder: Saurabh Mandal — full-stack builder + AI system architect. Built the entire from-scratch framework (2k+ lines, 9/9 demo gauntlet passing), plus 3 prior products (ASTRO, ALICE, SIA). Unconditional MSc Applied AI for Engineering offer from University of Greenwich [VERIFIED] — active research credibility.")
note("[ADVISOR NOTE: add 1-2 incubator-assigned mentors post-selection.]")

# 8. Traction / Demo
h1("8. Traction / Working Proof")
bullet("From-scratch nano transformer trained on CPU — 9/9 demos pass (text, code, vision, audio-listen, audio-gen, image, video, tools, quantize).")
bullet("Audio→language loop live (hear → describe → reply → speak back). VAE decoder gives real 256×256 images.")
bullet("LoRA fine-tune pipeline ready (Gemma 3 1B, Colab T4).")
bullet("Full investor doc suite: master report, appendices, TAM, financials, competitive, GTM, hiring, valuation [VERIFIED].")
p("Live demo URL: [LINK to demo video — REQUIRED]")

# 9. Use of Funds (the ask)
h1("9. Use of Funds — Ask ₹25,00,000 [FILL-IN]")
table_rows = [
    ("Fine-tuning Indic small models (GPU credits)", "₹9,00,000", "35%"),
    ("Team (founder stipend + 2 part-time)", "₹7,00,000", "28%"),
    ("Launch + pilot of revenue app (dev, store, ads)", "₹5,00,000", "20%"),
    ("Legal/company/compliance (SPICe+, DPDP)", "₹2,00,000", "8%"),
    ("Edge silicon eval kits + travel", "₹2,00,000", "8%"),
]
t = doc.add_table(rows=1+len(table_rows), cols=3); t.style = "Light Grid Accent 1"
for i, hh in enumerate(["Item", "Amount", "Share"]): t.rows[0].cells[i].text = hh
for r, row in enumerate(table_rows, start=1):
    for c, val in enumerate(row): t.rows[r].cells[c].text = val
doc.add_paragraph()
note("Adjust to portal cap (TIDE 2.0 up to ₹25L seed; some calls ₹50L with co-funding).")

# 10. Milestones (18 months)
h1("10. Milestones (18-month plan)")
table_rows = [
    ("Q1–Q2", "Import/INC-32 done · SIA Edge LoRA trained · demo v2 live", "₹0 extra (free tiers)"),
    ("Q3–Q4", "Revenue app launched · 5 pilots signed (enterprise/gov)", "₹5L from grant"),
    ("Q5–Q6", "100 pilot users · B2B white-label MoU · Series-extension ask", "₹10L from grant"),
]
t = doc.add_table(rows=1+len(table_rows), cols=3); t.style = "Light Grid Accent 1"
for i, hh in enumerate(["Phase", "Deliverables", "Funding"]): t.rows[0].cells[i].text = hh
for r, row in enumerate(table_rows, start=1):
    for c, val in enumerate(row): t.rows[r].cells[c].text = val
doc.add_paragraph()

# 11. Sustainability
h1("11. Sustainability Beyond Grant")
bullet("Pilot revenue (gov/enterprise) + B2B licensing create recurring cash.")
bullet("Founder co-investment + Amit (angel) seed at pilot stage [VERIFIED intent].")
bullet("Path to IndiaAI Mission / SISFS for scale round, if needed.")

# 12. Incubator partner
h1("12. Incubator Partner (required) [TODO]")
p("TIDE grants are routed through an incubator. Nearest: Bengaluru incubators (e.g. T-Hub network, IIT-Madras Incubator BLR hub, NSRCEL). Shortlist 2–3, pitch the demo, get LoI. Reason for choosing: AI/edge focus + DPIIT recognition.")

# 13. Checklist
h1("13. Pre-Submission Checklist")
for i in ["DP...", ""]:
    pass
for i in ["Startup India DPIIT recognition applied (free, fast)", "Company registered (SPICe+)", "Demo video 2 min linked", "Ca-F after incorporation", "Incubator LoI obtained", "PAN for company", "Founder ID/passport (Greenwich offer as education credential)"]:
    bullet("☐ " + i)

doc.save(OUT)
print(f"SAVED {OUT}")