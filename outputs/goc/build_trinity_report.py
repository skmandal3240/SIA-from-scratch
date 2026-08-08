#!/usr/bin/env python3
"""Build Project TRINITY Investor Report DOCX (30-page structure, SIA architecture filled).

Run: .venv/bin/python outputs/goc/build_trinity_report.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "Project_TRINITY_Investor_Report.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = NAVY
r = doc.add_paragraph("India's Sovereign AI Infrastructure Company\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Investor Report — Working Draft | August 2026\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(11)
r.runs[0].font.color.rgb = GREY
doc.add_paragraph("Prepared by SIA / Mandal Holdings. Illustrative ranges are planning estimates, not commitments.")

# ============ 1. EXEC SUMMARY ============
h1("1. Executive Summary")
p("Project TRINITY proposes building a complete AI technology stack for India across five strategic layers: Compute Infrastructure, Data Infrastructure, Foundation Models, AI Agents & AI Operating System, and AI Applications. The long-term vision is to create AI infrastructure and products serving enterprise, government, developers, and consumers — positioned not as a typical AI startup, but as India's sovereign AI infrastructure company.")
p("The delivery vehicle is SIA — a family of on-device and cloud AI models (Nano 0.5–1B to Cloud 70B+) sharing one tokenizer, one architecture, one tool-calling interface, one memory format, and one runtime. SIA is an AI operating system: private by design, built from scratch in India, embedded in every tier of hardware from wearables to data centers.")

# ============ 2. VISION / MISSION ============
h1("2. Vision, Mission & Why Now")
h2("Vision")
p("Build AI infrastructure, models, and applications that enable India to strengthen domestic AI capabilities while competing globally.")
h2("Mission")
p("Develop an integrated AI ecosystem covering infrastructure, models, developer platforms, and applications — the complete stack required for India to become an AI sovereign nation.")
h2("Why Now")
bullet("IndiaAI Mission: 34,000+ GPUs, 20 sovereign foundation models being funded")
bullet("DPDP 2023 makes on-device privacy a regulatory advantage")
bullet("Jevons paradox: cheaper AI drives exploding usage — cost leaders win")
bullet("Global race (US, China, EU) has no Indian full-stack player at sovereign scale")

# ============ 3. OPPORTUNITY ============
h1("3. India's AI Opportunity")
bullet("Limited sovereign AI infrastructure today — dependence on foreign cloud and models")
bullet("High compute costs; fragmented AI tooling; scarce domain-specific datasets")
bullet("India AI market growing >25%/yr; edge/on-device AI fastest segment")
bullet("900M+ users lack reliable connectivity — on-device AI is the only path to serve them")

# ============ 4. FIVE PILLARS ============
h1("4. AI's Five Pillars Framework")
table(["Pillar", "Scope"], [
    ("1. Compute", "GPU cloud, AI clusters, data centers, chips (CPU/GPU/TPU/NPU/photonic)"),
    ("2. Data", "Indian datasets, synthetic data, enterprise/government datasets, marketplace"),
    ("3. Models", "SIA family: Nano 0.5–1B / Edge 2–4B / Pro 8–14B / Cloud 70B+; vision, speech, multimodal"),
    ("4. Agents & OS", "AI OS, autonomous agents, coding agents, enterprise agents, robotics agents"),
    ("5. Applications", "Healthcare, education, agriculture, manufacturing, defence, government, consumer AI"),
])

# ============ 5. PILLAR 1 — COMPUTE ============
h1("5. Pillar 1 — Compute Infrastructure")
p("Infrastructure follows a staged build-vs-rent strategy:")
table(["Horizon", "Action"], [
    ("Years 1–2", "Rent GPU (Yotta, E2E, international) — focus on model development"),
    ("Years 2–4", "Build software infrastructure: distributed training, inference platform, MLOps, data pipelines"),
    ("Years 4–6", "Small private GPU cluster (~100–500 GPUs) for predictable workloads"),
    ("Years 6–10", "Hyperscale AI campus: thousands of GPUs, high-speed networking, dedicated substations, liquid cooling, peta-scale storage"),
])
p("Chip strategy: CPU for OS; GPU for training today; TPU-class tensor throughput; NPU for local inference (phones, laptops, robots, cars, IoT); optical/photonic for future data-center LLM workloads.")

# ============ 6. PILLAR 2 — DATA ============
h1("6. Pillar 2 — Data Infrastructure")
bullet("Indian-language + Indian-context datasets (the untapped moat)")
bullet("SIA Memory: consistent, personalized cross-device experience; compatible memory formats across tiers")
bullet("Synthetic data pipeline to scale scarce domain datasets")
bullet("Public + enterprise + government datasets; future data marketplace")

# ============ 7. PILLAR 3 — MODELS ============
h1("7. Pillar 3 — Foundation Models (SIA)")
p("SIA is an AI operating system — a family of models, not one giant model, all sharing the same stack:")
table(["Tier", "Size", "Target"], [
    ("SIA Nano", "0.5B – 1B", "Wearables, IoT, microcontrollers"),
    ("SIA Edge", "2B – 4B", "Phones, PCs, laptops"),
    ("SIA Pro", "8B – 14B", "Workstations, local GPUs"),
    ("SIA Cloud", "70B+", "Data centers — advanced reasoning"),
])
p("Shared across all tiers: tokenizer, architecture, tool-calling interface, memory formats, and a common runtime that synchronizes knowledge. A from-scratch nano transformer (text/audio/vision/code) is already trained on CPU; LoRA fine-tune pipeline is ready (Gemma 3 1B base).")

# ============ 8. PILLAR 4 — AGENTS & OS ============
h1("8. Pillar 4 — AI Agents & Operating System")
bullet("SIA Agent: autonomous task execution")
bullet("SIA Studio / SDK: developers embed SIA into apps and IoT products")
bullet("Common tool-calling interface (already demonstrated: calc/now/file tools via [[tool:name(args)]])")
bullet("Coding agents (Claude Code / Codex class), enterprise agents, workflow automation")

# ============ 9. PILLAR 5 — APPLICATIONS ============
h1("9. Pillar 5 — Applications & Ecosystem")
table(["Module", "Role"], [
    ("SIA Edge", "Local on-device AI model (delivery vehicle)"),
    ("SIA Core", "Central orchestration and synchronization"),
    ("SIA Cloud", "Advanced reasoning when local isn't enough"),
    ("SIA Studio / SDK", "Developer embed — apps and IoT"),
    ("SIA Memory", "Consistent personalized experience across devices"),
    ("SIA Vision", "Multimodal perception, recognition, synthesis"),
    ("SIA Voice", "Speech recognition + synthesis"),
    ("SIA Agent", "Autonomous task execution"),
])
p("Target verticals: healthcare, education, agriculture, manufacturing, defence, space, government, consumer AI.")

# ============ 10. ARCHITECTURE ============
h1("10. Technology Architecture")
p("The 5-Layer AI stack:")
table(["Layer", "Slice"], [
    ("L1", "Claude Code / Codex — coding agents"),
    ("L2", "Applications — All"),
    ("L3", "Models — foundation + SIA family"),
    ("L4", "Infrastructure — data centers"),
    ("L5", "Chip + Energy (nuclear path)"),
])
p("Outcomes: SIA Everywhere (models embedded in every device) and Root India / Build India (IP, jobs, chips, energy in India).")

# ============ 11. ROADMAP ============
h1("11. Product Roadmap (2026–2035)")
table(["Phase", "Objective", "Horizon"], [
    ("1", "Prototype & MVP (SIA Edge on-device companion)", "2026–2027"),
    ("2", "Foundation models (Nano/Edge trained, Pro LoRA)", "2027–2029"),
    ("3", "Enterprise platform (agents, SDK, memory)", "2029–2031"),
    ("4", "National AI infrastructure (cluster → campus)", "2031–2035"),
    ("5", "Global expansion", "2035+"),
])

# ============ 12. COMPETITIVE LANDSCAPE ============
h1("12. Competitive Landscape")
bullet("Global: OpenAI/Anthropic/Google (cloud-locked, English-first, no India sovereignty)")
bullet("China: sovereign stack exists — India has no equivalent")
bullet("India: no credible full-stack on-device + sovereign player today")
bullet("Moat: from-scratch stack, on-device privacy, Indian languages/context, silicon roadmap (DLI/PLI)")

# ============ 13. MARKET (TAM/SAM/SOM) ============
h1("13. Market Opportunity (TAM/SAM/SOM)")
bullet("TAM: global AI infrastructure + applications (multi-trillion)")
bullet("SAM: India AI market (enterprise + government + consumer) — $25B+ by 2030 est.")
bullet("SOM: on-device Indian-language AI (first-mover, no credible competitor)")
p("Note: detailed TAM/SAM/SOM model to be completed in Appendices (financial model).")

# ============ 14. BUSINESS MODEL ============
h1("14. Business Model")
bullet("AI API usage (SIA Cloud tier)")
bullet("Enterprise subscriptions (SIA Pro / agents)")
bullet("GPU cloud (post-Year 4 private cluster)")
bullet("Managed AI services + licensing")
bullet("Professional services / SDK embed fees")

# ============ 15. GTM ============
h1("15. Go-to-Market Strategy")
p("Sequenced: Developers → Startups → SMEs → Enterprises → Government. Developer-first via SIA Studio/SDK; government via sovereign AI narrative + grants.")

# ============ 16. FINANCIAL MODEL / FUNDING ============
h1("16. Financial Model & Funding Strategy")
table(["Stage", "Goal", "Illustrative Range"], [
    ("MVP", "Prototype", "₹2–10 Cr"),
    ("Growth", "Foundation models", "₹25–100 Cr"),
    ("Scale", "AI cloud", "₹100–500 Cr"),
    ("National", "Infrastructure", "₹500–2,000 Cr"),
    ("Global", "Expansion", "₹2,000 Cr+"),
])
p("Illustrative planning ranges only — actual funding depends on milestones, eligibility, investor appetite, and program requirements.")
h2("Government sources")
p("IndiaAI Mission, Startup India, MeitY (TIDE 2.0 / SAMRIDH), ANRF, DST (SISFS, NIDHI-PRAYAS), Digital India, Semiconductor Mission (DLI/PLI), IN-SPACe, state startup missions. Each has distinct eligibility, equity-vs-grant, timeline, and milestone requirements.")
h2("Private sources")
p("Angels, deep-tech VCs, AI-focused VCs, strategic investors, hyperscalers, semiconductor companies, sovereign wealth funds.")

# ============ 17. RISKS ============
h1("17. Risk Analysis")
bullet("Capital intensity (infrastructure) — mitigated by staged build-vs-rent")
bullet("Competition (hyperscalers) — differentiated by sovereignty + on-device privacy")
bullet("Technology execution — mitigated by working from-scratch stack + LoRA pipeline")
bullet("Regulatory changes — DPDP alignment is an advantage")
bullet("Talent acquisition — India dev talent pool + remote-first")

# ============ 18. AGI POSITION ============
h1("18. AGI Roadmap & Scientific Position")
p("Evidence-based framing only. Current AI capabilities are narrow; scientific limitations remain in reasoning, memory, compute, data, and robotics. Project TRINITY is designed to build the infrastructure, models, and ecosystem that position the company to contribute to increasingly capable AI systems over time. Whether AGI is achieved remains scientifically uncertain and cannot be guaranteed.")

# ============ 19. TEAM ============
h1("19. Team & Hiring Plan")
bullet("Founder: Saurabh Mandal — built the from-scratch SIA framework (tokenizer, transformer, multimodal, tools), 3 prior products (ASTRO, ALICE, SIA)")
bullet("Hiring plan: 1–2 junior engineers (Year 1), ML engineers + infra (Year 2–3), full team at enterprise phase")
p("Detailed hiring plan in Appendices.")

# ============ 20. ASK & USE OF FUNDS ============
h1("20. Investment Ask & Use of Funds")
p("Phased asks aligned to milestones above. Near-term ask: ₹2–10 Cr MVP phase — team, prototype, compute. Use of funds for the MVP phase:")
table(["Use", "Share"], [
    ("Team (founder + 1–2 engineers)", "~30%"),
    ("GPU compute (rented: Yotta/E2E/Colab)", "~35%"),
    ("Data + fine-tuning (Indian languages)", "~20%"),
    ("Legal, compliance, filings, SDK tooling", "~15%"),
])

# ============ 21. CLOSING ============
h1("21. Closing Vision & Call to Action")
p("We are building the complete AI technology stack required for India to become an AI sovereign nation — SIA, an AI operating system embedded in every device, with India as its root. Join us in building the infrastructure, models, and applications that keep India's data in India, India's intelligence in India, and India's AI future in Indian hands.")

# ============ APPENDIX: DUE DILIGENCE ============
h1("Appendix A — Investor Due Diligence (Key Q&A)")
h2("Compute — why build instead of rent?")
p("Ideal: rent early (Y1–2) to focus on models; build software infra (Y2–4); own GPU cluster only when predictable utilization justifies capex (Y4–6). Metrics: utilization %, $/GPU-hr, model-quality per ₹. Risk: capex timing. Evidence: staged roadmap above.")
h2("Why will enterprises switch?")
p("Sovereignty (data stays in India), on-device privacy (DPDP), Indian-language context, and cost at scale. Ideal: pilot with one enterprise + one government use case.")
h2("Government — why should taxpayers fund this?")
p("Jobs created, IP retained in India, export potential, national-security-relevant AI, rural impact, alignment with IndiaAI Mission and semiconductor roadmap. Must answer: how many jobs, how many patents, how much IP stays in India.")
h2("Valuation strategy")
p("Not startup-multiple on revenue alone — position as infrastructure: asset-backed milestones (GPU cluster, models, patents, government contracts) + revenue. Full model in Appendices.")

# ============ APPENDIX: GOVERNMENT EVAL ============
h1("Appendix B — Government Evaluation Criteria")
bullet("Alignment with IndiaAI Mission and Digital India")
bullet("Job creation (quantified)")
bullet("IP retention in India (patents, models, chips)")
bullet("Export potential and national security impact")
bullet("Rural/ESG impact")
bullet("Semiconductor roadmap alignment (DLI/PLI)")

# ============ APPENDIX: DUE DILIGENCE CHECKLIST ============
h1("Appendix C — Remaining Work (30-page completion)")
bullet("Detailed TAM/SAM/SOM model with sources")
bullet("Full financial model (3-statement, unit economics)")
bullet("Market sizing per vertical (health/edu/agri/defence/gov)")
bullet("Technical architecture diagrams (5-layer stack visuals)")
bullet("Hiring plan with costs")
bullet("Risk register (probability × impact)")
bullet("References list (policy, market, academic)")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")
