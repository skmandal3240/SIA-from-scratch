#!/usr/bin/env python3
"""Project TRINITY — Go-to-Market Strategy DOCX (standalone)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_GTM_Strategy.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Go-to-Market Strategy\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("August 2026 | SIA / Mandal Holdings | Working Draft\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
p("Placement: developer-first wedge, then government and enterprise. The GTM is deliberately sequenced so every phase funds and de-risks the next.")

# ============ 1. PHILOSOPHY ============
h1("1. GTM Philosophy")
bullet("Start with the developer wedge: SIA Studio/SDK embeds spread the product with zero sales cost.")
bullet("Convert trust into contracts: government + enterprise buy sovereignty before features.")
bullet("Phase-gated: each phase's revenue/grant funds the next step (matches the funding ladder).")
p("Sequencing: Developers → Startups → SMEs → Enterprises → Government.")

# ============ 2. CHANNEL MAP ============
h1("2. Channel Map")
table(["Seq", "Channel", "Mechanic", "Primary KPI"],
[
    ("1", "Developers (SIA Studio)", "Open-source + SDK + docs", "#SDK embeds, #downloads"),
    ("2", "Indian startups", "Freemium (Edge/API)", "Signups, MoM growth"),
    ("3", "SMEs", "Managed pilots (managed cloud)", "Pilot closure → paid"),
    ("4", "Enterprises", "Direct PoC + white-label", "PoC → 1st contract, LTV"),
    ("5", "Government", "Sovereign RFPs, grants", "Contract count, grant approved"),
])

# ============ 3. WEDGE & LAND-AND-EXPAND ============
h1("3. Wedge, Land & Expand")
h2("3.1 The wedge")
p("For developers: a free SDK that embeds SIA Edge into their app in minutes — local, private, Indic. For government: an on-device transliteration/dictation pilot in any department whose data must not leave India.")
h2("3.2 Land & expand")
p("ClirEntry: sign 3–5 enterprise pilots (Y1). Expand: one slot → full department → multi-year (SIA Pro + agents). Government: start with a district-level pilot → state minister-level contract.")
h2("3.3 Pricing (illustrative)")
table(["Product", "Price", "Notes"],
[
    ("SIA Edge app (consumer)", "Free + ₹99–299/mo premium", "Privacy as premium"),
    ("SIA API", "₹0.5–2 / 1K tokens", "Cloud tier"),
    ("SIA Pro (enterprise)", "₹5–20 lakh/yr", "White-label optional"),
    ("Government", "Milestone contracts", "Sovereign RFP route"),
])

# ============ 4. ACTIVITIES BY PHASE ============
h1("4. Activities by Phase")
h2("Phase 1 — Developer wedge (Y0–Y1)")
bullet("Ship open-source SDK; write 10 tutorials covering Hindi-EN code samples.")
bullet("Engineered meet-ups / hackathons in Bengaluru; engage on X/LinkedIn.")
bullet("Launch a 'privacy-first' page + DPDP compliance plain-english guide as free content.")
h2("Phase 2 — Startups & SMBs (Y1–Y2)")
bullet("Freemium API — measure signup → first paid conversion.")
bullet("Reference pilots; junior BD person; content for specific verticals (health, edu, agri).")
h2("Phase 3 — Enterprise (Y2–Y3)")
bullet("Direct sales + PoC account plan (target 3–5 accounts).")
bullet("White-label SDK for OEM/IoT (phones, kiosks, vehicles).")
h2("Phase 4 — Government (Y3+)")
bullet("Sovereign RFP tracking (IndiaAI, state digital missions); grants pipeline; demo video.")
bullet("Think-tank white papers on sovereignty + DPDP for decision-makers.")

# ============ 5. CONVERSION FUNNEL ============
h1("5. Conversion Funnel (indicative)")
table(["Stage", "Y1 target", "Y2 target"],
[
    ("Developer signups (SDK)", "5,000", "20,000"),
    ("Companies embedding", "50", "300"),
    ("Free → Paid (A/P)", "10%", "15%"),
    ("Paid accounts", "5 startups", "45"),
    ("Enterprise PoCs", "3–5", "10"),
    ("Government contracts", "0", "1–2"),
])
p("Movement: paid startups → insert PoCs → government contracts; every stage funds the next.")

# ============ 6. COMPETITIVE GTM STANCE ============
h1("6. Competitive GTM Stance")
bullet("We do NOT out-claim OpenAI. We own not the 'general' narrative — we own: privacy, Indic, sovereign, edge. That is a channel competitors cannot follow (structural).")
bullet("Against wrappers: we show the from-scratch stack — cost and IP advantage.")
bullet("Against foreign clouds: DPDP, data residency, price and latency at the edge.")

# ============ 7. KPIs & REVIEW ============
h1("7. KPIs & Cadence")
table(["Metric", "Ownership", "Cadence"],
      [("SDK embeds / devs", "Founder / community", "Weekly"),
       ("PoC pipeline value", "Founder", "Monthly"),
       ("Revenue & ARR", "Finance", "Monthly"),
       ("Grant pipeline", "Founder", "Fortnightly"),
       ("Gov contract funnel", "Founder + advisor", "Monthly")])
p("Review: weekly funnel; monthly pillar review; quarterly plan refresh; grant cycle = external trigger.")

# ============ 8. WHAT'S SKIPPED / NEXT ============
h1("8. What's Skipped (YAGNI at this stage)")
bullet("International expansion before India scale.")
bullet("Heavy brand/PR spend before product-market-fit.")
bullet("Hardware willship: SDK first; chip work guarded to DLI phase.")
p("Next practical step: 5 enterprise/government pilots + open-source SDK v0.1 + the documented GTM tracker.")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")