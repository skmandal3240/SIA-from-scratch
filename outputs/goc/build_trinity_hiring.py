#!/usr/bin/env python3
"""Project TRINITY — Hiring Plan DOCX (standalone)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_Hiring_Plan.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Team & Hiring Plan\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(16)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("August 2026 | SIA / Mandal Holdings | Working Draft\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
p("Costs are indicative annual ₹ lakh/crore. Geography: Bengaluru HQ, remote-first India. Founder-led until pilots convert.")

# ============ 1. TEAM PHILOSOPHY ============
h1("1. Team Philosophy")
bullet("Tiny senior core first — quality over headcount; every hire must unblock a roadmap item.")
bullet("1–2 junior engineers early keeps burn low and develops internal pipeline.")
bullet("Remote-first India (Bengaluru HQ anchor); ESOP pool target 10–15% from first institutional round.")
bullet("Contractor + grants for spike work; keep core team permanent.")

# ============ 2. FOUNDER TEAM ============
h1("2. Founder Team")
table(["Role", "Who", "Proof"],
[
    ("Founder / CEO + AI lead", "Saurabh Mandal", "From-scratch SIA stack; 3 prior products (ASTRO, ALICE, SIA)"),
    ("Advisors (target)", "TBD [TODO]", "Ex-IndiaAI/MeitY; ex-cloud; NLP/edge academic"),
])
p("Actively recruiting advisor pool to add credibility and RFP access.")

# ============ 3. HIRING PLAN BY PHASE ============
h1("3. Hiring Plan by Phase (costs indicative)")
h2("Phase 1 — MVP (Y1)")
table(["Role", "Headcount", "Annual cost (indicative)"],
[
    ("Founder (CEO/AI)", "1", "—"),
    ("ML/backend engineer", "1–2", "₹12–20L each"),
    ("Product/design (contract)", "1", "₹10–15L"),
    ("Ops/finance (part-time)", "0.5", "₹5–8L"),
])
h2("Phase 2 — Foundation models (Y2)")
table(["Role", "Headcount", "Annual cost (indicative)"],
[
    ("ML/LLM engineer", "2", "₹15–25L each"),
    ("Data engineer", "1", "₹12–18L"),
    ("Infra/DevOps (GPU)", "1", "₹15–20L"),
    ("Business development", "1", "₹10–15L + comp"),
])
h2("Phase 3 — Enterprise platform (Y3)")
table(["Role", "Headcount", "Annual cost (indicative)"],
[
    ("Platform engineer", "2–3", "₹20–30L each"),
    ("Agents/research eng", "2", "₹20–28L"),
    ("Enterprise AE/CS", "2", "₹15–25L + comp"),
    ("Product manager", "1", "₹20–30L"),
])
h2("Phase 4 — National infra (Y4+)")
table(["Role", "Headcount", "Annual cost (indicative)"],
[
    ("Cluster / infra team", "3–5", "₹25–40L each"),
    ("Research scientist", "2–3", "₹30–45L"),
    ("Field/ops + compliance", "2–3", "₹15–25L"),
    ("Supporting roles", "varies", "₹15–30L"),
])

# ============ 4. TOTAL HEADCOUNT & BURN ============
h1("4. Total Headcount & Burn")
table(["Year", "Headcount", "Annual comp. burn (indicative)"],
[
    ("Y1", "3", "₹25–40L"),
    ("Y2", "7", "₹60–80L"),
    ("Y3", "15", "₹1.2–1.5 Cr"),
    ("Y4", "30", "₹3 Cr+"),
])
p("Matching the financial model: comp equals ~60% of opex; compute next. ESOP pool carved at first institutional round.")

# ============ 5. KEY ROLES DESCRIPTIONS ============
h1("5. Key Role Descriptions")
h2("ML / LLM engineer (first hire)")
p("Own model training + LoRA fine-tunes; experience with transformers, PyTorch, HF ecosystem; comfortable on CPU + rented GPU workflows.")
h2("Data engineer")
p("Builds Indian-language datasets pipeline, synthetic data gen, SIA Memory formats.")
h2("Business development (Y2)")
p("Sells PoCs into enterprise + tracks government RFPs; prefers individuals with India gov experience (preferred advisory link).")
h2("Infra/DevOps (GPU)")
p("Runs rented-GPU workflows; later cluster ops: SLURM/horovod, storage, monitoring; security-minded (DPDP).")

# ============ 6. CULTURE ============
h1("6. Culture & Working Norms")
bullet("Remote-first, Bengaluru hub; async + weekly demo.")
bullet("High autonomy, written notes, no busy work.")
bullet("Everyone ships; docs are first-class.")
bullet("ESOP + mission alignment: 'Your private intelligence, on your device.'")

# ============ 7. HIRING TIMELINE ============
h1("7. Hiring Timeline (indicative)")
table(["Milestone", "Trigger"],
[
    ("First engineer", "Post-MVP closing (sign 1st pilot)"),
    ("Data engineer", "LoRA fine-tune begins"),
    ("BD hire", "3–5 PoCs signed"),
    ("Infra eng", "First GPU-training campaign"),
    ("Full platform team", "Scale round ₹100–500 Cr closes"),
])

# ============ 8. WHAT'S SKIPPED ============
h1("8. What's Skipped (YAGNI)")
bullet("No org-chart photo before 10+ people.")
bullet("No fractional leadership hires until funding justifies.")
bullet("Outsourced legal/CA as needed — not in-house until Y3.")
p("Next: definitive role write-ups + ESOP plan once incorporation closes.")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")