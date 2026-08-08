#!/usr/bin/env python3
"""Build standalone Appendix B — Government Evaluation Criteria DOCX for Project TRINITY.

Answers the 'Government Questions' from the TRINITY starter (why taxpayer money, jobs,
IP retention, patents, exports, national security, IndiaAI alignment, semiconductor
roadmap, rural impact, ESG) + a funding-source matrix.
Run: python3 outputs/goc/build_trinity_appendix_b.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "outputs" / "goc" / "TRINITY_Appendix_B_Government_Evaluation_Bengaluru.docx"

NAVY = RGBColor(0x0F, 0x1B, 0x2D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x55, 0x5F, 0x6E)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(18)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GOLD
        r.font.size = Pt(14)


def q(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = NAVY


def a(text):
    doc.add_paragraph(text)


def num(text):
    p = doc.add_paragraph()
    r = p.add_run("Evidence/answer: " + text)
    r.italic = True
    r.font.color.rgb = GREY


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    doc.add_paragraph()


# ============ COVER ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project TRINITY\n")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = NAVY
r = doc.add_paragraph("Appendix B — Government Evaluation Criteria\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(14)
r.runs[0].font.color.rgb = GOLD
r = doc.add_paragraph("Working draft | August 2026 | SIA / Mandal Holdings\n")
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
r.runs[0].font.size = Pt(10)
r.runs[0].font.color.rgb = GREY
doc.add_paragraph("Why taxpayer funding, what India gets back, and how Project TRINITY aligns with national AI, semiconductor, and rural-impact mandates. Figures are illustrative targets unless marked verified.")

# ============ 1. THE GOVERNMENT QUESTION ============
h1("1. The Core Government Question")
q("Why should taxpayers fund this?")
a("Because the alternative is strategic dependence: India currently relies on foreign cloud providers and foreign foundation models for its most sensitive AI workloads. Project TRINITY builds the sovereign alternative — infrastructure, models, and applications owned and operated in India. Taxpayer funding buys four public goods: national AI independence, domestic jobs, in-country IP, and infrastructure that can serve government, defence, and national-security workloads without data leaving the country. Every rupee returns hard assets (models, clusters, patents, trained engineers) that stay in India.")
num("Positioned as national infrastructure, not a startup subsidy — a deliberate, evidence-backed pitch.")

# ============ 2. DIRECT QUESTION & ANSWERS ============
h1("2. Government Questions & Answers")

q("How many jobs will you create?")
a("Phase 1 (MVP): 1–2 junior engineers. Phase 2 (models): 5–10 ML/infra roles. Phase 3 (enterprise platform): 15–30. Phase 4 (national infrastructure): 50–100+ direct, plus indirect ecosystem jobs (packaging, chip design, data annotation, service providers). Founders commit to hiring primarily within India, anchored in the Bengaluru hub, with jobs distributed across states to support national inclusion.")
num("Illustrative: 2 (Yr 1) → 10 (Yr 3) → 60+ (Yr 7). Verify against hiring plan.")

q("How much IP stays in India?")
a("100% of core IP: model architecture, training pipeline, tokenizer, datasets, and tool-calling stack are built from scratch in India (not wrappers around foreign APIs). Patents, trademarks, and model weights will be Indian-owned. No licensing-in of core models. Source code lives in Indian-owned repositories; the company is incorporated in India; nothing is assigned offshore.")
num("Target: 100% India-owned core IP; 0 foreign API dependency in the from-scratch stack.")

q("How many patents?")
a("Target: 2–4 Indian patent filings in years 1–2 covering the model architecture, multimodal encoders, and training methodology; 1 trademark application for 'SIA' (Classes 9 + 42). Trade secrets for datasets and fine-tuning methodology, protected via NDAs and IP-assignment contracts.")
num("Illustrative filing timeline: architecture patent Q1 2027; encoder/training patent Q3 2027.")

q("What is the export potential?")
a("Sovereign AI is an exportable Indian product: on-device models, SDKs, and the 'SIA OS' stack can be licensed to enterprises and governments in the Global South (Africa, Southeast Asia, South Asia) that share India's sovereignty and privacy requirements. The model family is a natural export once domestic maturity is proven — and exports earn foreign exchange, an explicit government evaluation criterion.")
num("Illustrative: exports target >25% of revenue by year 6–8.")

q("What is the national security impact?")
a("Defence, government, and critical-sector workloads run on foreign AI today, creating strategic dependence. SIA provides an on-device/sovereign alternative: models, data, and infrastructure under Indian control, auditable, DPDP-compliant, suitable for defence research, secure government workflows, and critical infrastructure. This is a direct contribution to national resilience, consistent with the Atmanirbhar Bharat direction.")
num("Use cases: secure government document processing, on-device analytics for field deployments, defence research assistance.")

# ============ 3. ALIGNMENT WITH NATIONAL PROGRAMS ============
h1("3. Alignment with National Programs")

q("Alignment with IndiaAI Mission?")
a("Direct. The IndiaAI Mission's 5 pillars (compute, data, models, talent, applications) map 1:1 to TRINITY's Pillars. SIA targets IndiaAI compute credits, interoperates with the AI Kosh dataset platform, and can submit under the foundation-model pillar alongside the 20 Asia-approved sovereign models. The staged rent→build compute strategy matches IndiaAI's phased subsidized-compute approach.")
num("Computed credits: eligible startups up to 40% reduced cost; 34,000+ GPUs in IndiaAI compute (context).")

q("Semiconductor roadmap?")
a("TRINITY's chip strategy (P4) aligns with the Semiconductor Mission and Design Linked Incentive: edge NPU for local inference, Dholera 28nm fab for tape-out, RISC-V domestic IP. The plan targets a first edge-AI test chip at 28nm Dholera trial production, leveraging DLI's ~50% design-cost coverage and the Make-in-India silicon mandate — an explicit semiconductor-roadmap contribution.")
num("Target: tape-out readiness assessment by Year 3; DLI application at Year 2–3.")

q("Rural impact?")
a("On-device AI (SIA Edge on low-end phones) works offline, so rural users without reliable connectivity get AI in Hindi and regional languages — healthcare, agriculture advisories, education, government-service navigation. TRINITY's applications pillar explicitly targets agriculture and education for rural India, directly serving the digital-inclusion mandate. Rural deployment pilots can run from the Bengaluru corporate base into adjoining Karnataka and across states, keeping field presence close to the home state as it scales.")
num("Target: 1 rural use case pilot in Karnataka (agriculture advisory) by Year 2, then replicate to other states.")

q("ESG impact?")
a("Environmental: on-device inference cuts data-center energy per query; future energy path aligns with India's nuclear buildout (100 GW goal) — cheaper, cleaner power for AI. Social: privacy + inclusion (offline, Indic languages). Governance: DPDP-aligned, audited, transparent. TRINITY ties its three-stage nuclear energy plan (PHWR → PFBR → thorium) to its infra roadmap; the energy cost curve is an explicit cost advantage.")
num("Energy math: nuclear 100 GW goal vs data-center build-out; on-device cuts $/query by ~10× vs cloud inference (illustrative).")

# ============ 4. FUNDING SOURCE MATRIX ============
h1("4. Government Funding Sources Matrix")
table(["Program", "Ministry/Dept", "Illustrative Max", "Equity vs Grant", "Alignment"],
[
    ("IndiaAI Mission", "MeitY", "Compute credits + model support", "Grant/subsidy", "Compute, models, data"),
    ("TIDE 2.0", "MeitY", "up to ₹50L", "Grant", "Incubation-linked, tech"),
    ("SAMRIDH", "MeitY", "up to ₹40L", "Grant/equity (varies)", "Startup accelerator"),
    ("SISFS", "DST", "up to ₹50L", "Grant", "Seed support"),
    ("NIDHI-PRAYAS", "DST", "₹10L", "Grant", "Early prototyping"),
    ("DLI (Design Linked Incentive)", "MeitY", "~50% design cost", "Incentive", "Edge-ASIC design"),
    ("PLI (Make in India)", "DPIIT", "varies", "Incentive", "Hardware manufacturing"),
    ("Karnataka State Startup/IT policy", "Govt of Karnataka", "varies", "Grant/equity", "Bengaluru-registered entity"),
])
num("Verify each program's current call window, eligibility, and milestone reporting before applying — details in 04_TIDE2_DOSSIER.md and 02_GRANTS_TRACKER.md.")

# ============ 5. GOVERNMENT EVALUATION CHECKLIST ============
h1("5. Likely Government Evaluation Checklist")
for item in [
    "IndiaAI Mission and Digital India alignment — 1:1 pillar mapping shown",
    "Job creation (quantified) — see §2.1 trajectory and hiring plan",
    "IP retention in India — 100% India-owned core IP target",
    "Export potential — >25% revenue export target by Y6–8",
    "National security impact — sovereign/on-device positioning",
    "Rural and ESG impact — rural pilot + offline/Indic access + clean-energy linkage",
    "Semiconductor roadmap — Dholera 28nm, DLI/PLI alignment",
    "Financial viability, milestone-based — phase-gated ask, no hockey-stick",
]:
    bullet = doc.add_paragraph(style="List Bullet")
    bullet.add_run(item)

# ============ 6. SCORECARD SELF-ASSESSMENT ============
h1("6. Self-Assessment Scorecard (1–5)")
table(["Criteria", "Score (1–5)", "Gap to close"],
[
    ("National security value", "4", "Defence-specific use case write-up"),
    ("IP retention in India", "5", "Patent filings scheduled"),
    ("IndiaAI/MeitY alignment", "4", "Formal application + DPIIT number"),
    ("Jobs + talent", "3", "Hiring plan with costs"),
    ("Rural/ESG", "3", "Rural pilot defined"),
    ("Export potential", "3", "Global-South GTM sheet"),
    ("Semiconductor roadmap", "3", "DLI/PLI application"),
    ("Financial/sustainability", "3", "Full 3-statement model"),
])
doc.add_paragraph("Self-assessment: median 3.5/5 — application-ready once the 'gap to close' items land. Those are exactly the items in the investor report's Appendix C.")

doc.save(OUT)
print(f"SAVED {OUT} ({len(doc.paragraphs)} paragraphs)")